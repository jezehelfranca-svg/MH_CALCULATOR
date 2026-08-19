# -*- coding: utf-8 -*-
from pathlib import Path
import base64, tempfile, math

HYUNDAI_BLUE = '004EA2'
DARK_BLUE = '17365D'
LIGHT_BLUE = 'D9EAF7'
LIGHT_GRAY = 'F2F2F2'
PALE_GREEN = 'E2F0D9'
PALE_YELLOW = 'FFF2CC'


def _num(v):
    try:
        if v in (None, '', '-'): return 0.0
        return float(str(v).replace(',', ''))
    except Exception:
        return 0.0


def _fmt(v, d=0):
    try:
        return f"{float(v):,.{d}f}"
    except Exception:
        return str(v)


def _pct(a, b):
    a=_num(a); b=_num(b)
    return a / b * 100 if b else 0


def _set_cell_shading(cell, fill):
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill)
        tcPr.append(shd)
    except Exception:
        pass


def _set_cell_text(cell, text, bold=False, color=None, size=8, align='left'):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else: p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _style_table(table, header_fill=DARK_BLUE, header_font='FFFFFF'):
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = 'Malgun Gothic'
            if i == 0:
                _set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = __import__('docx').shared.RGBColor.from_string(header_font)


def _add_heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    p.style = f'Heading {min(level,3)}'
    r = p.add_run(text)
    r.font.name = 'Malgun Gothic'
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r.font.size = Pt(15 if level == 1 else 12)
    return p


def _add_note_box(doc, title, lines, fill='EDF4FB'):
    table = doc.add_table(rows=1, cols=1)
    _style_table(table, header_fill=fill, header_font='000000')
    cell = table.cell(0,0)
    _set_cell_shading(cell, fill)
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.font.name = 'Malgun Gothic'; r.bold = True
    for line in lines:
        p = cell.add_paragraph('• ' + str(line))
        for rr in p.runs: rr.font.name = 'Malgun Gothic'
    doc.add_paragraph('')


def _chart_images(data, workdir):
    """Create chart images. If matplotlib is unavailable, return empty list."""
    images=[]
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        try:
            plt.rcParams['font.family'] = ['Malgun Gothic', 'AppleGothic', 'NanumGothic', 'DejaVu Sans']
            plt.rcParams['axes.unicode_minus'] = False
        except Exception:
            pass
        total=data.get('total',{})
        ci=data.get('ci',{}); tel=data.get('tel',{})
        # Pie chart internal/external
        vals=[_num(total.get('internal')), _num(total.get('external'))]
        if sum(vals) > 0:
            fig, ax = plt.subplots(figsize=(5.2,3.2), dpi=160)
            ax.pie(vals, labels=['내부 M/H','외주 M/H'], autopct='%1.1f%%', startangle=90, colors=['#004EA2','#7FB3D5'])
            ax.set_title('내부/외주 M-H 구성')
            fn=Path(workdir)/'chart_internal_external.png'
            fig.tight_layout(); fig.savefig(fn, transparent=False); plt.close(fig)
            images.append(('내부/외주 M-H 구성', fn))
        # Bar chart CI/TEL
        fig, ax = plt.subplots(figsize=(6.4,3.4), dpi=160)
        labels=['C&I','Telecom']
        internal=[_num(ci.get('internal')), _num(tel.get('internal'))]
        external=[_num(ci.get('external')), _num(tel.get('external'))]
        x=range(len(labels))
        ax.bar([i-0.18 for i in x], internal, width=0.36, label='내부', color='#004EA2')
        ax.bar([i+0.18 for i in x], external, width=0.36, label='외주', color='#70AD47')
        ax.set_xticks(list(x)); ax.set_xticklabels(labels)
        ax.set_ylabel('M/H')
        ax.set_title('분야별 내부/외주 M-H')
        ax.legend()
        ax.grid(axis='y', alpha=0.25)
        fn=Path(workdir)/'chart_part_breakdown.png'
        fig.tight_layout(); fig.savefig(fn); plt.close(fig)
        images.append(('분야별 내부/외주 M-H', fn))
    except Exception:
        return []
    return images


def create_word_report(data, path, logo_b64=None):
    """Create a Korean executive-style FEED M/H Word report from calculator data."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT

    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.55)
    sec.right_margin = Inches(0.55)

    styles = doc.styles
    styles['Normal'].font.name = 'Malgun Gothic'
    styles['Normal'].font.size = Pt(9)

    # Cover/header block
    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = True
    left, right = header.rows[0].cells
    _set_cell_shading(left, 'FFFFFF'); _set_cell_shading(right, 'FFFFFF')
    left.text=''
    if logo_b64:
        try:
            img_path = path.parent / '_hyundai_logo_tmp.png'
            img_path.write_bytes(base64.b64decode(logo_b64))
            left.paragraphs[0].add_run().add_picture(str(img_path), width=Inches(1.55))
            try: img_path.unlink()
            except Exception: pass
        except Exception:
            run = left.paragraphs[0].add_run('HYUNDAI\nENGINEERING')
            run.font.bold=True; run.font.color.rgb=RGBColor.from_string(HYUNDAI_BLUE); run.font.size=Pt(14)
    else:
        run = left.paragraphs[0].add_run('HYUNDAI\nENGINEERING')
        run.font.bold=True; run.font.color.rgb=RGBColor.from_string(HYUNDAI_BLUE); run.font.size=Pt(14)
    right.text=''
    p=right.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    r=p.add_run('FEED M/H 산출 보고서')
    r.font.name='Malgun Gothic'; r.font.size=Pt(18); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK_BLUE)
    p=right.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run('Control & Instrumentation / Telecommunication').font.size=Pt(9)
    doc.add_paragraph('')

    total=data.get('total',{})
    ci=data.get('ci',{}); tel=data.get('tel',{})
    project=data.get('project') or '-'
    base=data.get('base_mh',161); months=data.get('months',4)

    meta = doc.add_table(rows=1, cols=4)
    meta.rows[0].cells[0].text='Project'
    meta.rows[0].cells[1].text=str(project)
    meta.rows[0].cells[2].text='작성 기준'
    meta.rows[0].cells[3].text='FEED M/H Calculator V0.1'
    row=meta.add_row().cells
    row[0].text='외주최소화'; row[1].text=str(data.get('external_min','No'))
    row[2].text='Base / Duration'; row[3].text=f'{_fmt(base)} M/H per M/M / {_fmt(months)}개월'
    row=meta.add_row().cells
    row[0].text='외주 단가'; row[1].text=f"단종 {_fmt(data.get('rate_short'))} 원 / 종합 {_fmt(data.get('rate_comp'))} 원"
    row[2].text='보고 범위'; row[3].text='C&I + Telecom FEED Engineering M/H'
    _style_table(meta)
    doc.add_paragraph('')

    _add_heading(doc, '1. Executive Summary', 1)
    sm = doc.add_table(rows=1, cols=6)
    headers=['구분','내부 M/H','외주 M/H','Total M/H','Total M/M','평균 투입인원']
    for i,h in enumerate(headers): sm.rows[0].cells[i].text=h
    for label, obj in [('C&I',ci),('Telecom',tel),('Total',total)]:
        row=sm.add_row().cells
        row[0].text=label
        row[1].text=_fmt(obj.get('internal'))
        row[2].text=_fmt(obj.get('external'))
        row[3].text=_fmt(obj.get('total'))
        row[4].text=_fmt(obj.get('mm'),1)
        row[5].text=_fmt(obj.get('avg'),1)
    _style_table(sm)
    doc.add_paragraph('')

    # cost table
    cost = doc.add_table(rows=1, cols=4)
    for i,h in enumerate(['비용 항목','산출 기준','예상 비용','비고']): cost.rows[0].cells[i].text=h
    for label, key, rate_key in [('외주-단종 예가','short_cost','rate_short'),('외주-종합 예가','comp_cost','rate_comp')]:
        row=cost.add_row().cells
        row[0].text=label
        row[1].text=f"외주 M/H {_fmt(total.get('external'))} × 단가 {_fmt(data.get(rate_key))}"
        row[2].text=f"{_fmt(data.get(key))} 원"
        row[3].text='VAT 및 계약 조건 별도 검토'
    _style_table(cost)
    doc.add_paragraph('')

    # Analysis callouts
    total_mh=_num(total.get('total'))
    in_pct=_pct(total.get('internal'), total_mh)
    ex_pct=_pct(total.get('external'), total_mh)
    bigger='C&I' if _num(ci.get('total')) >= _num(tel.get('total')) else 'Telecom'
    _add_note_box(doc, '주요 분석', [
        f"총 산출 M/H는 {_fmt(total_mh)} M/H이며, 내부 {_fmt(in_pct,1)}%, 외주 {_fmt(ex_pct,1)}% 구성입니다.",
        f"분야별로는 {bigger} 비중이 더 크며, 상세 산출 근거는 Activity별 산출 상세표에 정리했습니다.",
        "외주최소화 적용 시 내부/외주 배분비가 변경되므로, 외주 예가와 투입계획을 함께 검토해야 합니다.",
        "본 보고서는 FEED 단계의 기술자료 검토, Data Sheet, Drawing, MTO/Inform, 3D/SPI/CER 등 주요 역무 기준으로 작성되었습니다."
    ])

    # Charts
    with tempfile.TemporaryDirectory() as tmp:
        charts=_chart_images(data, tmp)
        if charts:
            _add_heading(doc, '2. M/H 구성 차트', 1)
            for title, img in charts:
                p=doc.add_paragraph()
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER
                p.add_run(title).bold=True
                doc.add_picture(str(img), width=Inches(5.8))
                doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('')

    _add_heading(doc, '3. 분야별 Activity 산출 상세', 1)
    for title, part in [('3.1 C&I 상세', ci), ('3.2 Telecom 상세', tel)]:
        _add_heading(doc, title, 2)
        rows=part.get('rows', [])[:45]
        tbl=doc.add_table(rows=1, cols=8)
        headers=['Code','Activity','Unit','Qty','난이도','내부 M/H','외주 M/H','Total']
        for i,h in enumerate(headers): tbl.rows[0].cells[i].text=h
        for r in rows:
            row=tbl.add_row().cells
            vals=[r.get('code',''), r.get('activity',''), r.get('unit',''), _fmt(r.get('qty'),1), r.get('diff',''), _fmt(r.get('hec')), _fmt(r.get('ext')), _fmt(r.get('total'))]
            for i,v in enumerate(vals): row[i].text=str(v)
        _style_table(tbl)
        doc.add_paragraph('')

    _add_heading(doc, '4. 산출 기준 및 적용 원칙', 1)
    _add_note_box(doc, '산출 기준 요약', [
        '수량 입력값과 Project Condition 선택값을 기준으로 Activity별 내부/외주 Unit M/H를 적용합니다.',
        '난이도는 Project 표준 난이도, Specification/Data Sheet/Drawing 조건, SPI 적용 여부 등을 기준으로 산정합니다.',
        'M/M은 Total M/H를 Base M/H로 나누어 산출하며, 평균 투입인원은 M/M을 설계기간으로 나누어 산정합니다.',
        '외주 예가는 외주 M/H와 외주 적용단가를 곱한 예산 검토용 값입니다.'
    ], fill='F8FBFE')

    std=data.get('std',{})
    for part_key, title in [('ci','C&I 대표 산출 기준'),('tel','Telecom 대표 산출 기준')]:
        entries=list((std.get(part_key) or {}).values())[:18]
        if not entries: continue
        _add_heading(doc, title, 2)
        tbl=doc.add_table(rows=1, cols=6)
        for i,h in enumerate(['Activity','Unit','내부 기준','외주 기준','Guide','비고']): tbl.rows[0].cells[i].text=h
        for x in entries:
            row=tbl.add_row().cells
            row[0].text=str(x.get('activity',''))
            row[1].text=str(x.get('unit',''))
            row[2].text='/'.join([f"{k}:{_fmt(v,2)}" for k,v in (x.get('int') or {}).items() if _num(v)!=0]) or '-'
            row[3].text='/'.join([f"{k}:{_fmt(v,2)}" for k,v in (x.get('ext') or {}).items() if _num(v)!=0]) or '-'
            row[4].text=str(x.get('guide',''))[:420]
            row[5].text='난이도별 Unit M/H'
        _style_table(tbl)
        doc.add_paragraph('')

    _add_heading(doc, '5. 검토 의견', 1)
    _add_note_box(doc, '검토 및 후속 조치', [
        '대형 Quantity 항목은 입력 수량의 신뢰성이 전체 M/H에 직접적인 영향을 주므로, 견적 전 최종 수량 확인이 필요합니다.',
        '외주 비중이 큰 역무는 산출조건, 외주 Scope, 제출물 수준을 명확히 정의하여 예가 산정 편차를 줄이는 것이 필요합니다.',
        'Project Condition 선택값이 난이도와 산출 기준에 반영되므로, ITB 및 발주처 요구사항 검토 결과와 일치하는지 확인해야 합니다.',
        '본 보고서는 FEED 단계의 M/H 산출 검토용이며, EPC 전환 또는 상세설계 착수 시 산출기준과 수량을 재검토해야 합니다.'
    ], fill='FFF8E6')

    # footer
    sec = doc.sections[0]
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Hyundai Engineering FEED M/H Calculator | Generated Report')
    run.font.name='Malgun Gothic'; run.font.size=Pt(8); run.font.color.rgb=RGBColor.from_string('666666')

    doc.save(str(path))
    return str(path)
