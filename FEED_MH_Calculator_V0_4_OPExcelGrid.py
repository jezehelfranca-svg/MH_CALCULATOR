
# -*- coding: utf-8 -*-
import json,csv,math,os,subprocess,sys,base64,tempfile
from pathlib import Path
try:
    import report_utils
except Exception:
    class report_utils:
        @staticmethod
        def create_word_report(data, path, logo_b64=None):
            try:
                from docx import Document
            except Exception as e:
                raise RuntimeError('Word Report 생성을 위해 python-docx 모듈이 필요합니다: '+str(e))
            from pathlib import Path
            doc=Document()
            doc.add_heading('FEED M/H Report', 0)
            doc.add_paragraph('HYUNDAI ENGINEERING')
            summary=data.get('total',{}) if isinstance(data,dict) else {}
            tbl=doc.add_table(rows=1, cols=2)
            tbl.rows[0].cells[0].text='Item'; tbl.rows[0].cells[1].text='Value'
            for k,v in [('Project',data.get('project','-')),('Part',data.get('part','-')),('Case',data.get('case','-')),('External Min',data.get('external_min','-')),('Internal M/H',summary.get('internal',0)),('External M/H',summary.get('external',0)),('Total M/H',summary.get('total',0))]:
                row=tbl.add_row().cells; row[0].text=str(k); row[1].text=str(v)
            doc.add_heading('Guide / Help', 1)
            doc.add_paragraph('Summary, Input, Output, OP, 산출기준 탭을 확인하여 M/H 산출 결과를 검토합니다.')
            Path(path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(path))
import tkinter as tk
from tkinter import ttk,filedialog,messagebox
DATA=json.loads('{"inputs": [{"part": "ci", "code": "CTRL_PROJECT", "section": "0. 프로젝트 공통", "major": "Project", "middle": "", "detail": "PROJECT", "value": "", "unit": "", "type": "text", "options": [], "note": "프로젝트명", "id": 0}, {"part": "ci", "code": "CTRL_EXTERNAL_MIN", "section": "0. 프로젝트 공통", "major": "외주최소화", "middle": "", "detail": "외주최소화 적용 유무", "value": "No", "unit": "", "type": "select", "options": [{"value": "No", "text": "No 일반 산출 기준"}, {"value": "Yes", "text": "Yes 외주최소화 75/25"}], "note": "Yes 선택 시 내부 75%, 외부 25% 적용", "id": 1}, {"part": "ci", "code": "A01_QTY", "baseCode": "A01", "section": "Key Input (수량)", "major": "Key Input (수량)", "middle": "Control System", "detail": "Total ICSS (DCS+ESD+F&G)+MMS IO 수량", "value": 8820, "unit": "Point", "type": "number", "options": [], "note": "Actual수량만 기입", "id": 2}, {"part": "ci", "code": "A02_QTY", "baseCode": "A02", "section": "A. 수량 입력", "major": "", "middle": "Instruments", "detail": "Total Instrument 수량 (In-line Instrument 포함)", "value": 6462, "unit": "Set", "type": "number", "options": [], "note": "F&G Inst. & In-line Inst. 포함", "id": 3}, {"part": "ci", "code": "A03_QTY", "baseCode": "A03", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "In-line Instrument (수량 Valve & In-line Instruments)", "value": 0, "unit": "Set", "type": "number", "options": [], "note": "수량 집계 가능 시 직접 입력, 불가 시 0을 입력 (0을 입력 시, A15 물량은 A02의 25%로 자동 반영 됨)", "id": 4}, {"part": "ci", "code": "A04_QTY", "baseCode": "A04", "section": "GENERAL", "major": "GENERAL", "middle": "ITB Review & Clarification", "detail": "Meeting & Coordination", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 5}, {"part": "ci", "code": "A05_QTY", "baseCode": "A05", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Review & Interface work.", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 6}, {"part": "ci", "code": "A06_QTY", "baseCode": "A06", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "발주처 협의 (Technical Clarification)", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 7}, {"part": "ci", "code": "A07_QTY", "baseCode": "A07", "section": "Specification", "major": "Specification", "middle": "Technical Specification", "detail": "Instrumentation & Automation Design Criteria", "value": 1, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 8}, {"part": "ci", "code": "A08_QTY", "baseCode": "A08", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Control System", "value": 1, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 9}, {"part": "ci", "code": "A09_QTY", "baseCode": "A09", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Packaged Instrument (Metering, Analzyer, TGS & etc.)", "value": 2, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 10}, {"part": "ci", "code": "A10_QTY", "baseCode": "A10", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Field Instrument", "value": 3, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 11}, {"part": "ci", "code": "A11_QTY", "baseCode": "A11", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Others", "value": 1, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 12}, {"part": "ci", "code": "A12_QTY", "baseCode": "A12", "section": "Calculation", "major": "Calculation", "middle": "Calculation", "detail": "Cv (Valve Coefficient) Calculation for Sizing", "value": 485, "unit": "Set", "type": "number", "options": [], "note": "Quantity의 경우 수량 집계 가능 시 직접 입력, 불가 시 Inline 계기의 30%로 반영 (수식 적용), Calculation 미수행 시 0으로 입력", "id": 13}, {"part": "ci", "code": "A13_QTY", "baseCode": "A13", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "PSV Orificice Calculation for Sizing", "value": 323, "unit": "Set", "type": "number", "options": [], "note": "Quantity의 경우 수량 집계 가능 시 직접 입력, 불가 시 Inline 계기의 20%로 반영 (수식 적용), Calculation 미수행 시 0으로 입력", "id": 14}, {"part": "ci", "code": "A14_QTY", "baseCode": "A14", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Flow Element Calculation", "value": 242, "unit": "Set", "type": "number", "options": [], "note": "Quantity의 경우 수량 집계 가능 시 직접 입력, 불가 시 Inline 계기의 15%로 반영 (수식 적용), Calculation 미수행 시 0으로 입력", "id": 15}, {"part": "ci", "code": "A15_QTY", "baseCode": "A15", "section": "구매관련", "major": "구매관련", "middle": "Data Sheets & Summary List", "detail": "Data Sheets (Valve & In-line Instruments)", "value": 1616, "unit": "Ea", "type": "number", "options": [], "note": "수량 입력", "id": 16}, {"part": "ci", "code": "A16_QTY", "baseCode": "A16", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Summary List (Other Instruments)", "value": 4846, "unit": "Ea", "type": "number", "options": [], "note": "수량 입력", "id": 17}, {"part": "ci", "code": "A17_QTY", "baseCode": "A17", "section": "A. 수량 입력", "major": "", "middle": "MR & TBE", "detail": "Control System", "value": 1, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 18}, {"part": "ci", "code": "A18_QTY", "baseCode": "A18", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Packaged Instrument", "value": 3, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 19}, {"part": "ci", "code": "A19_QTY", "baseCode": "A19", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Valves", "value": 7, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 20}, {"part": "ci", "code": "A20_QTY", "baseCode": "A20", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Field Instrument", "value": 10, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 21}, {"part": "ci", "code": "A21_QTY", "baseCode": "A21", "section": "DRAWING", "major": "DRAWING", "middle": "Control System Architecture", "detail": "Control System Configuration", "value": 2, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 22}, {"part": "ci", "code": "A22_QTY", "baseCode": "A22", "section": "A. 수량 입력", "major": "", "middle": "Layout Drawing", "detail": "Instrument Cable Routing Layout", "value": 18, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 23}, {"part": "ci", "code": "A23_QTY", "baseCode": "A23", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Cable Tray, Trench, Duct Bank Layout", "value": 18, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 24}, {"part": "ci", "code": "A24_QTY", "baseCode": "A24", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Junction Box & Wiring Layout \\n(Sketched Drawing)", "value": 18, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 25}, {"part": "ci", "code": "A25_QTY", "baseCode": "A25", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Control Room / PIB Equipment Arrangement Layout", "value": 3, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 26}, {"part": "ci", "code": "A26_QTY", "baseCode": "A26", "section": "A. 수량 입력", "major": "", "middle": "Summary / List", "detail": "Instrument Index and I/O List", "value": 6462, "unit": "장", "type": "number", "options": [], "note": "외주 최소화 적용시 \'내부 업무 중 외주-종합 수행 비율 (%)\'은 0을 입력 할 것.", "id": 27}, {"part": "ci", "code": "A27_QTY", "baseCode": "A27", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Instrument Cable Schedule", "value": 1293, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 28}, {"part": "ci", "code": "A28_QTY", "baseCode": "A28", "section": "A. 수량 입력", "major": "", "middle": "Detail Drawing", "detail": "Typical Hook-up drawings", "value": 81, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 29}, {"part": "ci", "code": "A29_QTY", "baseCode": "A29", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Instrument Typical Installation Details", "value": 50, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 30}, {"part": "ci", "code": "A30_QTY", "baseCode": "A30", "section": "A. 수량 입력", "major": "", "middle": "3\'rd Party System Interface Design & Cable Block Diagram", "detail": "Major 3rd Party System (Comp. / Boiler / Metering / MMS / Analyzer System / IRP (MCC) / FACP 등)", "value": 12, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 31}, {"part": "ci", "code": "A31_QTY", "baseCode": "A31", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Minor 3rd Party System (Pump / Chem. Injection / TGS / PAGA / ECMS / HVAC / NOVAC 등)", "value": 20, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 32}, {"part": "ci", "code": "A32_QTY", "baseCode": "A32", "section": "MTO & 타부서 INFORM", "major": "MTO & 타부서 INFORM", "middle": "MTO", "detail": "BM Take-off & BOQ작성", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 33}, {"part": "ci", "code": "A33_QTY", "baseCode": "A33", "section": "A. 수량 입력", "major": "", "middle": "타부서 Inform", "detail": "타부서 Inform 용 자료작성", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 34}, {"part": "ci", "code": "A34_QTY", "baseCode": "A34", "section": "OTHERS", "major": "OTHERS", "middle": "3D Model", "detail": "3D Modeling", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 35}, {"part": "ci", "code": "A35_QTY", "baseCode": "A35", "section": "A. 수량 입력", "major": "", "middle": "SPI", "detail": "SPI Seed DB", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 36}, {"part": "ci", "code": "A36_QTY", "baseCode": "A36", "section": "A. 수량 입력", "major": "", "middle": "CER", "detail": "Cold Eye Review", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 37}, {"part": "ci", "code": "C01_SEL", "baseCode": "C01", "section": "B. Project Condition", "major": "Project 표준 난이도\\n(General)", "middle": "FEED 난이도", "detail": "FEED 난이도", "value": "2)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) EPC 수행을 염두에 둔 정확도와 비교적 높은 난이도 및 상세가 요구되는 FEED 사업의 경우. (e.g. Aramco, Minor Revamping 포함 등)"}, {"value": "2)", "text": "2) EPC 수행을 염두에 둔 정확도가 요구되는 일반적인 FEED 사업의 경우. (e.g. 대부분의 FEED, 제안형 사업)"}, {"value": "3)", "text": "3) FEED 참여는 하되, EPC 수행 예정이 없는 FEED 사업의 경우"}], "note": "선택값에 따라 난이도 자동 산정", "id": 38}, {"part": "ci", "code": "C01_DIFF", "baseCode": "C01", "section": "B. Project Condition", "major": "Project 표준 난이도\\n(General)", "middle": "FEED 난이도", "detail": "FEED 난이도 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 39}, {"part": "ci", "code": "C04_SEL", "baseCode": "C04", "section": "B. Project Condition", "major": "Technical Specification 난이도", "middle": "Instrumentation & Automation Design Criteria", "detail": "Instrumentation & Automation Design Criteria", "value": "2)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) New Generation"}, {"value": "2)", "text": "2) HEC Standard Specification 적용"}], "note": "Project 표준 난이도와 연동하여 난이도 결정", "id": 40}, {"part": "ci", "code": "C04_DIFF", "baseCode": "C04", "section": "B. Project Condition", "major": "Technical Specification 난이도", "middle": "Instrumentation & Automation Design Criteria", "detail": "Instrumentation & Automation Design Criteria 난이도", "value": "하", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 41}, {"part": "ci", "code": "C06_SEL", "baseCode": "C06", "section": "B. Project Condition", "major": "", "middle": "Specification for Control System", "detail": "Specification for Control System", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) New Generation"}, {"value": "2)", "text": "2) HEC Standard Specification 적용"}], "note": "Project 표준 난이도와 연동하여 난이도 결정", "id": 42}, {"part": "ci", "code": "C06_DIFF", "baseCode": "C06", "section": "B. Project Condition", "major": "", "middle": "Specification for Control System", "detail": "Specification for Control System 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 43}, {"part": "ci", "code": "C08_SEL", "baseCode": "C08", "section": "B. Project Condition", "major": "", "middle": "Specification for Packaged Instrument (Metering, Analzyer, TGS & etc.)", "detail": "Specification for Packaged Instrument (Metering, Analzyer, TGS & etc.)", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) New Generation"}, {"value": "2)", "text": "2) HEC Standard Specification 적용"}], "note": "Project 표준 난이도와 연동하여 난이도 결정", "id": 44}, {"part": "ci", "code": "C08_DIFF", "baseCode": "C08", "section": "B. Project Condition", "major": "", "middle": "Specification for Packaged Instrument (Metering, Analzyer, TGS & etc.)", "detail": "Specification for Packaged Instrument (Metering, Analzyer, TGS & etc.) 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 45}, {"part": "ci", "code": "C10_SEL", "baseCode": "C10", "section": "B. Project Condition", "major": "", "middle": "Specification for Field Instrument", "detail": "Specification for Field Instrument", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) New Generation"}, {"value": "2)", "text": "2) HEC Standard Specification 적용"}], "note": "Project 표준 난이도와 연동하여 난이도 결정", "id": 46}, {"part": "ci", "code": "C10_DIFF", "baseCode": "C10", "section": "B. Project Condition", "major": "", "middle": "Specification for Field Instrument", "detail": "Specification for Field Instrument 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 47}, {"part": "ci", "code": "C12_SEL", "baseCode": "C12", "section": "B. Project Condition", "major": "", "middle": "Specification for Others", "detail": "Specification for Others", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) New Generation"}, {"value": "2)", "text": "2) HEC Standard Specification 적용"}], "note": "Project 표준 난이도와 연동하여 난이도 결정", "id": 48}, {"part": "ci", "code": "C12_DIFF", "baseCode": "C12", "section": "B. Project Condition", "major": "", "middle": "Specification for Others", "detail": "Specification for Others 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 49}, {"part": "ci", "code": "C14_SEL", "baseCode": "C14", "section": "B. Project Condition", "major": "Instrument Data Sheet 난이도", "middle": "Data Sheet (Valve & In-line Instruments)", "detail": "Data Sheet (Valve & In-line Instruments)", "value": "2)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) SPI 적용-내부"}, {"value": "2)", "text": "2) SPI 적용-외부"}, {"value": "3)", "text": "3) SPI 미적용"}], "note": "선택값에 따라 난이도 자동 산정", "id": 50}, {"part": "ci", "code": "C14_DIFF", "baseCode": "C14", "section": "B. Project Condition", "major": "Instrument Data Sheet 난이도", "middle": "Data Sheet (Valve & In-line Instruments)", "detail": "Data Sheet (Valve & In-line Instruments) 난이도", "value": "SPI-외부", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 51}, {"part": "ci", "code": "C19_SEL", "baseCode": "C19", "section": "B. Project Condition", "major": "Drawing 난이도", "middle": "Control System Configuration", "detail": "Control System Configuration", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) ITB에 해당 도면이 없어 Concept 확정 필요"}, {"value": "2)", "text": "2) ITB에 해당 도면이 포함되어 있으나 상당부분 Upgrade가 필요"}], "note": "선택값에 따라 난이도 자동 산정", "id": 52}, {"part": "ci", "code": "C19_DIFF", "baseCode": "C19", "section": "B. Project Condition", "major": "Drawing 난이도", "middle": "Control System Configuration", "detail": "Control System Configuration 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 53}, {"part": "ci", "code": "C21_SEL", "baseCode": "C21", "section": "B. Project Condition", "major": "", "middle": "Control Room / PIB Equipment Arrangement layout", "detail": "Control Room / PIB Equipment Arrangement layout", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) ITB에 해당 도면이 없어 도면 작성 필요"}, {"value": "2)", "text": "2) ITB에 해당 도면이 포함되어 있을 경우"}], "note": "선택값에 따라 난이도 자동 산정", "id": 54}, {"part": "ci", "code": "C21_DIFF", "baseCode": "C21", "section": "B. Project Condition", "major": "", "middle": "Control Room / PIB Equipment Arrangement layout", "detail": "Control Room / PIB Equipment Arrangement layout 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 55}, {"part": "ci", "code": "C23_SEL", "baseCode": "C23", "section": "B. Project Condition", "major": "", "middle": "Instrument Index", "detail": "Instrument Index", "value": "3)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) SPI 적용-내부"}, {"value": "2)", "text": "2) SPI 적용-외부"}, {"value": "3)", "text": "3) SPI 미적용"}], "note": "선택값에 따라 난이도 자동 산정", "id": 56}, {"part": "ci", "code": "C23_DIFF", "baseCode": "C23", "section": "B. Project Condition", "major": "", "middle": "Instrument Index", "detail": "Instrument Index 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 57}, {"part": "ci", "code": "C29_SEL", "baseCode": "C29", "section": "B. Project Condition", "major": "OTHERS 난이도", "middle": "3D Modeling", "detail": "3D Modeling", "value": "2)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) 일반적인 FEED Modleing & Review 역무 대비 축소"}, {"value": "2)", "text": "2) 일반적인 FEED Modleing & Review 역무"}, {"value": "3)", "text": "3) 일반적인 FEED Modleing & Review 역무 대비 확대"}], "note": "선택값에 따라 난이도 자동 산정", "id": 58}, {"part": "ci", "code": "C29_DIFF", "baseCode": "C29", "section": "B. Project Condition", "major": "OTHERS 난이도", "middle": "3D Modeling", "detail": "3D Modeling 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 59}, {"part": "ci", "code": "C35_VALUE", "baseCode": "C35", "section": "B. Project Condition", "major": "M/H 산출 General", "middle": "Base M/H", "detail": "Base M/H (per 1M/M)", "value": 161, "unit": "M/H", "type": "number", "options": [], "note": "PM팀에서 접수", "id": 60}, {"part": "ci", "code": "C36_VALUE", "baseCode": "C36", "section": "B. Project Condition", "major": "", "middle": "설계기간", "detail": "총 설계기간 (개월)", "value": 4, "unit": "개월", "type": "number", "options": [], "note": "프로젝트 Schedule 확인", "id": 61}, {"part": "ci", "code": "C37_VALUE", "baseCode": "C37", "section": "B. Project Condition", "major": "", "middle": "외주 적용단가", "detail": "외주-단종 적용단가 (원)", "value": 32000, "unit": "원", "type": "number", "options": [], "note": "PM팀에서 접수", "id": 62}, {"part": "ci", "code": "C38_VALUE", "baseCode": "C38", "section": "B. Project Condition", "major": "", "middle": "", "detail": "외주-종합 적용단가 (원)", "value": 39000, "unit": "원", "type": "number", "options": [], "note": "PM팀에서 접수", "id": 63}, {"part": "tel", "code": "CTRL_PROJECT", "section": "0. 프로젝트 공통", "major": "Project", "middle": "", "detail": "PROJECT", "value": "", "unit": "", "type": "text", "options": [], "note": "프로젝트명", "id": 64}, {"part": "tel", "code": "CTRL_EXTERNAL_MIN", "section": "0. 프로젝트 공통", "major": "외주최소화", "middle": "", "detail": "외주최소화 적용 유무", "value": "No", "unit": "", "type": "select", "options": [{"value": "No", "text": "No 일반 산출 기준"}, {"value": "Yes", "text": "Yes 외주최소화 75/25"}], "note": "Yes 선택 시 내부 75%, 외부 25% 적용", "id": 65}, {"part": "tel", "code": "A01_QTY", "baseCode": "A01", "section": "Key Input (수량)", "major": "Key Input (수량)", "middle": "System Q\'ty", "detail": "Telecom.&System Q\'ty", "value": 7, "unit": "Ea", "type": "number", "options": [], "note": "수량 입력", "id": 66}, {"part": "tel", "code": "A02_QTY", "baseCode": "A02", "section": "A. 수량 입력", "major": "", "middle": "Site Size", "detail": "m x m", "value": 187500, "unit": "M^2", "type": "number", "options": [], "note": "수량 입력", "id": 67}, {"part": "tel", "code": "A03_QTY", "baseCode": "A03", "section": "A. 수량 입력", "major": "", "middle": "Building Q\'ty", "detail": "Building Q\'ty", "value": 7, "unit": "Ea", "type": "number", "options": [], "note": "if 2 stories or more than 70m, 2ea", "id": 68}, {"part": "tel", "code": "A04_QTY", "baseCode": "A04", "section": "GENERAL", "major": "GENERAL", "middle": "ITB Review & Clarification", "detail": "Meeting & Coordination", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 69}, {"part": "tel", "code": "A05_QTY", "baseCode": "A05", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "ITB Study", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 70}, {"part": "tel", "code": "A06_QTY", "baseCode": "A06", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Technical Clarification", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 71}, {"part": "tel", "code": "A07_QTY", "baseCode": "A07", "section": "구매관련", "major": "구매관련", "middle": "MR & TBE", "detail": "Telecom.&Security System", "value": 2, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 72}, {"part": "tel", "code": "A08_QTY", "baseCode": "A08", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Telecom.&Security Bulk Material", "value": 1, "unit": "Set", "type": "number", "options": [], "note": "수량 입력", "id": 73}, {"part": "tel", "code": "A09_QTY", "baseCode": "A09", "section": "System Architecture", "major": "System Architecture", "middle": "System Architecture", "detail": "Telecom.&Security Specification", "value": 7, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 74}, {"part": "tel", "code": "A10_QTY", "baseCode": "A10", "section": "DRAWING", "major": "DRAWING", "middle": "Block Diagram", "detail": "Telecom.&Security Block Diagram", "value": 7, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 75}, {"part": "tel", "code": "A11_QTY", "baseCode": "A11", "section": "A. 수량 입력", "major": "", "middle": "Layout Drawing", "detail": "Telecom.&Security Cable Routing layout", "value": 11.166666666666668, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 76}, {"part": "tel", "code": "A12_QTY", "baseCode": "A12", "section": "A. 수량 입력", "major": "", "middle": "", "detail": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "value": 5.583333333333334, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 77}, {"part": "tel", "code": "A13_QTY", "baseCode": "A13", "section": "A. 수량 입력", "major": "", "middle": "Summary / List", "detail": "Telecom.&Security Cable Schedule", "value": 18.75, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 78}, {"part": "tel", "code": "A14_QTY", "baseCode": "A14", "section": "A. 수량 입력", "major": "", "middle": "Detail Drawing", "detail": "Telecom.&Security Typical Installation Details", "value": 7, "unit": "장", "type": "number", "options": [], "note": "수량 입력", "id": 79}, {"part": "tel", "code": "A15_QTY", "baseCode": "A15", "section": "MTO & 타부서 INFORM", "major": "MTO & 타부서 INFORM", "middle": "MTO", "detail": "BM Take-off & BOQ작성", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 80}, {"part": "tel", "code": "A16_QTY", "baseCode": "A16", "section": "A. 수량 입력", "major": "", "middle": "타부서 Inform", "detail": "타부서 Inform 용 자료작성", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 81}, {"part": "tel", "code": "A17_QTY", "baseCode": "A17", "section": "A. 수량 입력", "major": "", "middle": "CER", "detail": "Cold Eye Review", "value": 1, "unit": "Lot", "type": "number", "options": [], "note": "수량 입력", "id": 82}, {"part": "tel", "code": "A18_QTY", "baseCode": "A18", "section": "OTHERS", "major": "OTHERS", "middle": "3D Model", "detail": "3D Modeling", "value": 1, "unit": "Ea", "type": "number", "options": [], "note": "수량 입력", "id": 83}, {"part": "tel", "code": "A19_QTY", "baseCode": "A19", "section": "A. 수량 입력", "major": "", "middle": "SPI", "detail": "Study & Calculation", "value": 1, "unit": "Ea", "type": "number", "options": [], "note": "수량 입력", "id": 84}, {"part": "tel", "code": "A20_QTY", "baseCode": "A20", "section": "A. 수량 입력", "major": "", "middle": "Tie-in", "detail": "Telecom. Tie-in List", "value": 1, "unit": "Ea", "type": "number", "options": [], "note": "수량 입력", "id": 85}, {"part": "tel", "code": "C01_SEL", "baseCode": "C01", "section": "B. Project Condition", "major": "Project 표준 난이도\\n(General)", "middle": "ITB Requirements", "detail": "ITB Requirements", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) ITB가 불분명하여 확인할 사항이 많은 경우"}, {"value": "2)", "text": "2) ITB 자료가 충분한 경우"}], "note": "선택값에 따라 난이도 자동 산정", "id": 86}, {"part": "tel", "code": "C01_DIFF", "baseCode": "C01", "section": "B. Project Condition", "major": "Project 표준 난이도\\n(General)", "middle": "ITB Requirements", "detail": "ITB Requirements 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 87}, {"part": "tel", "code": "C03_SEL", "baseCode": "C03", "section": "B. Project Condition", "major": "System 난이도", "middle": "System Block Diagram & Specification", "detail": "System Block Diagram & Specification", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) ITB에 없어 백지상태에서 작성 필요"}, {"value": "2)", "text": "2) ITB 자료가 충분한 경우"}], "note": "선택값에 따라 난이도 자동 산정", "id": 88}, {"part": "tel", "code": "C03_DIFF", "baseCode": "C03", "section": "B. Project Condition", "major": "System 난이도", "middle": "System Block Diagram & Specification", "detail": "System Block Diagram & Specification 난이도", "value": "중", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 89}, {"part": "tel", "code": "C05_SEL", "baseCode": "C05", "section": "B. Project Condition", "major": "Drawing 난이도", "middle": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "detail": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "value": "1)", "unit": "", "type": "select", "options": [{"value": "1)", "text": "1) 사우디, 쿠웨이트, UAE 사업"}, {"value": "2)", "text": "2) 그 외"}], "note": "선택값에 따라 난이도 자동 산정", "id": 90}, {"part": "tel", "code": "C05_DIFF", "baseCode": "C05", "section": "B. Project Condition", "major": "Drawing 난이도", "middle": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "detail": "Telecom.&Security Cable Tray, Trench, Ductbank Layout 난이도", "value": "상", "unit": "", "type": "computed", "options": [], "note": "자동계산 결과", "id": 91}, {"part": "tel", "code": "C09_VALUE", "baseCode": "C09", "section": "B. Project Condition", "major": "M/H 산출 General", "middle": "Base M/H", "detail": "Base M/H (per 1M/M)", "value": 161, "unit": "M/H", "type": "number", "options": [], "note": "PM팀에서 접수", "id": 92}, {"part": "tel", "code": "C10_VALUE", "baseCode": "C10", "section": "B. Project Condition", "major": "", "middle": "설계기간", "detail": "총 설계기간 (개월)", "value": 4, "unit": "개월", "type": "number", "options": [], "note": "프로젝트 Schedule 확인", "id": 93}, {"part": "tel", "code": "C11_VALUE", "baseCode": "C11", "section": "B. Project Condition", "major": "", "middle": "외주 적용단가", "detail": "외주-단종 적용단가 (원)", "value": 32000, "unit": "원", "type": "number", "options": [], "note": "PM팀에서 접수", "id": 94}, {"part": "tel", "code": "C12_VALUE", "baseCode": "C12", "section": "B. Project Condition", "major": "", "middle": "", "detail": "외주-종합 적용단가 (원)", "value": 39000, "unit": "원", "type": "number", "options": [], "note": "PM팀에서 접수", "id": 95}], "outputs": [{"part": "ci", "code": "A04", "activity": "Meeting & Coordination", "unit": "Lot", "qty": 1.0, "baseIntUnit": 150.0, "baseExtUnit": 0.0, "std": {"activity": "Meeting & Coordination", "unit": "Lot", "int": {"SPI": 0.0, "상": 220.0, "중": 150.0, "하": 80.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정 기준\\n- Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A05", "activity": "Review & Interface work.", "unit": "Lot", "qty": 1.0, "baseIntUnit": 100.0, "baseExtUnit": 200.0, "std": {"activity": "Review & Interface work.", "unit": "Lot", "int": {"SPI": 0.0, "상": 150.0, "중": 100.0, "하": 50.0}, "ext": {"SPI": 0.0, "상": 250.0, "중": 200.0, "하": 100.0}, "guide": ""}}, {"part": "ci", "code": "A06", "activity": "발주처 협의 (Technical Clarification)", "unit": "Lot", "qty": 1.0, "baseIntUnit": 100.0, "baseExtUnit": 0.0, "std": {"activity": "발주처 협의 (Technical Clarification)", "unit": "Lot", "int": {"SPI": 0.0, "상": 150.0, "중": 100.0, "하": 50.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A07", "activity": "Instrumentation & Automation Design Criteria", "unit": "Set", "qty": 1.0, "baseIntUnit": 80.0, "baseExtUnit": 0.0, "std": {"activity": "Instrumentation & Automation Design Criteria", "unit": "Set", "int": {"SPI": 0.0, "상": 200.0, "중": 160.0, "하": 80.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n - 상 : New Generation 및 Project 표준 난이도 \'상\'.\\n - 중 : New Generation 및 Project 표준 난이도 \'중\' / \'하\'.\\n - 하 : HEC Standard Specification 적용."}}, {"part": "ci", "code": "A08", "activity": "Control System", "unit": "Set", "qty": 1.0, "baseIntUnit": 120.0, "baseExtUnit": 0.0, "std": {"activity": "Control System", "unit": "Set", "int": {"SPI": 0.0, "상": 36.0, "중": 30.0, "하": 24.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n- Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A09", "activity": "Packaged Instrument (Metering, Analzyer, TGS & etc.)", "unit": "Set", "qty": 2.0, "baseIntUnit": 100.0, "baseExtUnit": 0.0, "std": {"activity": "Packaged Instrument (Metering, Analzyer, TGS & etc.)", "unit": "Set", "int": {"SPI": 0.0, "상": 130.0, "중": 100.0, "하": 40.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A10", "activity": "Field Instrument", "unit": "Set", "qty": 3.0, "baseIntUnit": 60.0, "baseExtUnit": 0.0, "std": {"activity": "Field Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 12.0, "중": 10.0, "하": 8.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A11", "activity": "Others", "unit": "Set", "qty": 1.0, "baseIntUnit": 60.0, "baseExtUnit": 0.0, "std": {"activity": "Others", "unit": "Set", "int": {"SPI": 0.0, "상": 80.0, "중": 60.0, "하": 20.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A12", "activity": "Cv (Valve Coefficient) Calculation for Sizing", "unit": "Set", "qty": 485.0, "baseIntUnit": 0.4, "baseExtUnit": 0.0, "std": {"activity": "Cv (Valve Coefficient) Calculation for Sizing", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 0.4, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 별도 성과품을 제출하거나, 직접 수행하는 경우만 M/H 반영.\\n2) 난이도는 일괄 \\"중\\" 적용."}}, {"part": "ci", "code": "A13", "activity": "PSV Orificice Calculation for Sizing", "unit": "Set", "qty": 323.0, "baseIntUnit": 0.4, "baseExtUnit": 0.0, "std": {"activity": "PSV Orificice Calculation for Sizing", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 0.4, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 별도 성과품을 제출하거나, 직접 수행하는 경우만 M/H 반영.\\n2) 난이도는 일괄 \\"중\\" 적용."}}, {"part": "ci", "code": "A14", "activity": "Flow Element Calculation", "unit": "Set", "qty": 242.0, "baseIntUnit": 0.4, "baseExtUnit": 0.0, "std": {"activity": "Flow Element Calculation", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 0.4, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 별도 성과품을 제출하거나, 직접 수행하는 경우만 M/H 반영.\\n2) 난이도는 일괄 \\"중\\" 적용."}}, {"part": "ci", "code": "A15", "activity": "Data Sheets (Valve & In-line Instruments)", "unit": "Ea", "qty": 1616.0, "baseIntUnit": 0.0, "baseExtUnit": 0.3, "std": null}, {"part": "ci", "code": "A16", "activity": "Summary List (Other Instruments)", "unit": "Ea", "qty": 4846.0, "baseIntUnit": 0.05, "baseExtUnit": 0.0, "std": {"activity": "Summary List (Other Instruments)", "unit": "Ea", "int": {"SPI": 0.0, "상": 0.06, "중": 0.05, "하": 0.05}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A17", "activity": "Control System", "unit": "Set", "qty": 1.0, "baseIntUnit": 30.0, "baseExtUnit": 0.0, "std": {"activity": "Control System", "unit": "Set", "int": {"SPI": 0.0, "상": 36.0, "중": 30.0, "하": 24.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n- Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A18", "activity": "Packaged Instrument", "unit": "Set", "qty": 3.0, "baseIntUnit": 20.0, "baseExtUnit": 0.0, "std": {"activity": "Packaged Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 24.0, "중": 20.0, "하": 16.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A19", "activity": "Valves", "unit": "Set", "qty": 7.0, "baseIntUnit": 15.0, "baseExtUnit": 0.0, "std": {"activity": "Valves", "unit": "Set", "int": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A20", "activity": "Field Instrument", "unit": "Set", "qty": 10.0, "baseIntUnit": 10.0, "baseExtUnit": 0.0, "std": {"activity": "Field Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 12.0, "중": 10.0, "하": 8.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A17", "activity": "Control System", "unit": "Set", "qty": 1.0, "baseIntUnit": 30.0, "baseExtUnit": 0.0, "std": {"activity": "Control System", "unit": "Set", "int": {"SPI": 0.0, "상": 36.0, "중": 30.0, "하": 24.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n- Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A18", "activity": "Packaged Instrument", "unit": "Set", "qty": 3.0, "baseIntUnit": 20.0, "baseExtUnit": 0.0, "std": {"activity": "Packaged Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 24.0, "중": 20.0, "하": 16.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A19", "activity": "Valves", "unit": "Set", "qty": 7.0, "baseIntUnit": 15.0, "baseExtUnit": 0.0, "std": {"activity": "Valves", "unit": "Set", "int": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A20", "activity": "Field Instrument", "unit": "Set", "qty": 10.0, "baseIntUnit": 10.0, "baseExtUnit": 0.0, "std": {"activity": "Field Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 12.0, "중": 10.0, "하": 8.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A21", "activity": "Control System Configuration", "unit": "A3", "qty": 2.0, "baseIntUnit": 40.0, "baseExtUnit": 4.0, "std": {"activity": "Control System Configuration", "unit": "A3", "int": {"SPI": 0.0, "상": 50.0, "중": 40.0, "하": 30.0}, "ext": {"SPI": 0.0, "상": 4.8, "중": 4.0, "하": 3.2}, "guide": "1) Draft only, design은 HEC Engineer가 수행\\n2) 난이도 산정기준\\n - 상 : ITB에 해당 도면이 없어 Concept 확정 필요 및 Project 표준 난이도 \'상\'\\n - 중 : ITB에 해당 도면이 없어 Concept 확정 필요 및 Project 표준 난이도 \'중\' / \'하\'\\n - 하 : ITB에 해당 도면이 포함되어 있으나 상당부분 Upgrade가 필요"}}, {"part": "ci", "code": "A22", "activity": "Instrument Cable Routing layout", "unit": "A1", "qty": 18.0, "baseIntUnit": 4.0, "baseExtUnit": 35.0, "std": {"activity": "Instrument Cable Routing layout", "unit": "A1", "int": {"SPI": 0.0, "상": 5.0, "중": 4.0, "하": 3.0}, "ext": {"SPI": 0.0, "상": 42.0, "중": 35.0, "하": 28.0}, "guide": "1) JB, LCP위치 표현\\n2) Cable Tray Route 및 Typical Tray Section 표현\\n3) Cable Trench, Ductbank Route 및 Typical Trench, Ductbank Section 표현\\n4) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A23", "activity": "Instrument Cable Tray, Trench, Ductbank Layout", "unit": "A3", "qty": 18.0, "baseIntUnit": 2.5, "baseExtUnit": 15.0, "std": {"activity": "Instrument Cable Tray, Trench, Ductbank Layout", "unit": "A3", "int": {"SPI": 0.0, "상": 3.0, "중": 2.5, "하": 2.0}, "ext": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "guide": "1) Tray Size, Routing, Layer, Main Support 표현\\n2) Trench Detail, Size, Routing 표현 \\n3) Duct Bank Size, Location 표현\\n4) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A24", "activity": "JB & Wiring Layout (Sketched Drawing)", "unit": "A3", "qty": 18.0, "baseIntUnit": 2.5, "baseExtUnit": 15.0, "std": {"activity": "JB & Wiring Layout (Sketched Drawing)", "unit": "A3", "int": {"SPI": 0.0, "상": 3.0, "중": 2.5, "하": 2.0}, "ext": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "guide": "1) Multi Cable & Tray Routing 표현\\n2) 계기, JB (PKG JB포함), LCP, LCS 위치 표현 \\n3) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A25", "activity": "Control Room / PIB Equipment Arrangement layout", "unit": "A3", "qty": 3.0, "baseIntUnit": 4.0, "baseExtUnit": 30.0, "std": {"activity": "Control Room / PIB Equipment Arrangement layout", "unit": "A3", "int": {"SPI": 0.0, "상": 5.0, "중": 4.0, "하": 3.0}, "ext": {"SPI": 0.0, "상": 36.0, "중": 30.0, "하": 24.0}, "guide": "1) Cabinet Type, Size, Scope (Package, Electric 등) 명기\\n2) 난이도 산정기준\\n - 상 : ITB에 해당 도면이 없어 도면 작성 필요 및 Project 표준 난이도 \'상\'\\n - 중 : ITB에 해당 도면이 없어 도면 작성 필요 및 Project 표준 난이도 \'중\' / \'하\'\\n - 하 : ITB에 해당 도면이 포함되어 있을 경우"}}, {"part": "ci", "code": "A26", "activity": "Instrument Index and I/O List", "unit": "Ea", "qty": 6462.0, "baseIntUnit": 0.3, "baseExtUnit": 0.15, "std": {"activity": "Instrument Index and I/O List", "unit": "Ea", "int": {"SPI": 0.5, "상": 0.36, "중": 0.3, "하": 0.24}, "ext": {"SPI": 0.3, "상": 0.18, "중": 0.15, "하": 0.12}, "guide": "1) Tag no., Inst. Type, P&ID no. \\n2) Piping Class, Line Size,  Location, System, I/O Type, etc.\\n3) 난이도는 Project 표준 난이도를 적용. 단, SPI입력 시 \'SPI\'"}}, {"part": "ci", "code": "A27", "activity": "Instrument Cable Schedule", "unit": "Ea", "qty": 1293.0, "baseIntUnit": 0.05, "baseExtUnit": 0.1, "std": {"activity": "Instrument Cable Schedule", "unit": "Ea", "int": {"SPI": 0.0, "상": 0.06, "중": 0.05, "하": 0.04}, "ext": {"SPI": 0.0, "상": 0.12, "중": 0.1, "하": 0.08}, "guide": "1) Cable no, From/To, Cable code, \\n2) Cable gland spec. & q\'ty\\n3) JB & Wiring Layout에 의한 Multi cable length적용\\n4) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A28", "activity": "Typical Hook-up drawings", "unit": "A3", "qty": 81.0, "baseIntUnit": 0.3, "baseExtUnit": 3.0, "std": {"activity": "Typical Hook-up drawings", "unit": "A3", "int": {"SPI": 0.0, "상": 0.36, "중": 0.3, "하": 0.24}, "ext": {"SPI": 0.0, "상": 3.6, "중": 3.0, "하": 2.6}, "guide": "1) Tag 별 Type Assign Index (Excel) 작성 필요\\n2) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A29", "activity": "Instrument Typical Installation Details", "unit": "A3", "qty": 50.0, "baseIntUnit": 0.2, "baseExtUnit": 2.0, "std": {"activity": "Instrument Typical Installation Details", "unit": "A3", "int": {"SPI": 0.0, "상": 0.24, "중": 0.2, "하": 0.16}, "ext": {"SPI": 0.0, "상": 2.4, "중": 2.0, "하": 1.6}, "guide": "1) Tag 별 Type Assign Index (Excel) 작성 필요\\n2) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A30", "activity": "- Major 3rd Party System Interface", "unit": "Set", "qty": 12.0, "baseIntUnit": 5.0, "baseExtUnit": 5.0, "std": {"activity": "- Major 3rd Party System Interface", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 5.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 5.0, "하": 0.0}, "guide": "1) Skid JB , LCP, LCS, Loosed Instrument informs\\n2) HEC Scope Cable에 대한 Cable informs\\n3) Vendor Data 및 Reference Project 자료 검토"}}, {"part": "ci", "code": "A31", "activity": "- Minor 3rd Party System Interface", "unit": "Set", "qty": 20.0, "baseIntUnit": 2.0, "baseExtUnit": 2.0, "std": {"activity": "- Minor 3rd Party System Interface", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 2.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 2.0, "하": 0.0}, "guide": ""}}, {"part": "ci", "code": "A32", "activity": "BM Take-off & BOQ작성", "unit": "Lot", "qty": 1.0, "baseIntUnit": 120.0, "baseExtUnit": 600.0, "std": {"activity": "BM Take-off & BOQ작성", "unit": "Lot", "int": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 최대 3차 MTO까지 산출\\n2) Cable은 Cable Schedule 작성 후 산출\\n3) Ladder Tray & Fitting은 실 도면을 기준으로 산출\\n4) Perporate Tray는 Cable Route Sketch 작성 후 산출\\n5) Cable Gland는 Cable schedule에서 산출\\n6) JB, Air Pipe & Fittings는 BM Sketch 작성 후 산출\\n7) Tubing, Hook-up Valve, Fitting, Gasket & Bolt는 Hook-up물량 산출표 (Excel)를 작성하여 산출"}}, {"part": "ci", "code": "A33", "activity": "타부서 Inform 용 자료작성", "unit": "Lot", "qty": 1.0, "baseIntUnit": 100.0, "baseExtUnit": 190.0, "std": {"activity": "타부서 Inform용 자료작성", "unit": "Lot", "int": {"SPI": 0.0, "상": 120.0, "중": 100.0, "하": 80.0}, "ext": {"SPI": 0.0, "상": 200.0, "중": 190.0, "하": 180.0}, "guide": "1) Cable Tray & Duct Loading Data\\n2) Cable Trench, Ductbank & Foundation Data\\n3) Electrical Load List\\n4) Heat Dissipation Data\\n5) Heat Tracing Instrument List (ET, ST, CT)\\n6) Field Grounding Inform\\n7) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A34", "activity": "3D Modeling", "unit": "Lot", "qty": 1.0, "baseIntUnit": 30.0, "baseExtUnit": 250.0, "std": {"activity": "3D Modeling", "unit": "Lot", "int": {"SPI": 0.0, "상": 40.0, "중": 30.0, "하": 20.0}, "ext": {"SPI": 0.0, "상": 300.0, "중": 250.0, "하": 200.0}, "guide": "1) 난이도는 Modeling & Review 역무 Volume에 따라 난이도를 적용"}}, {"part": "ci", "code": "A35", "activity": "SPI SEED DB", "unit": "Lot", "qty": 1.0, "baseIntUnit": 80.0, "baseExtUnit": 0.0, "std": {"activity": "SPI SEED DB", "unit": "Lot", "int": {"SPI": 0.0, "상": 100.0, "중": 80.0, "하": 60.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도는 Project 표준 난이도를 적용"}}, {"part": "ci", "code": "A36", "activity": "Cold Eye Review", "unit": "Lot", "qty": 1.0, "baseIntUnit": 20.0, "baseExtUnit": 0.0, "std": {"activity": "Cold Eye Review", "unit": "Lot", "int": {"SPI": 0.0, "상": 30.0, "중": 20.0, "하": 15.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도는 Project 표준 난이도를 적용"}}, {"part": "tel", "code": "A04", "activity": "Meeting & Coordination", "unit": "Lot", "qty": 1.0, "baseIntUnit": 60.0, "baseExtUnit": 0.0, "std": {"activity": "Meeting & Coordination", "unit": "Lot", "int": {"상": 0.0, "중": 60.0, "하": 30.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정 기준\\n - 중 : ITB가 불분명하여 확인할 사항이 많은 경우\\n - 하 : ITB 자료가 충분한 경우"}}, {"part": "tel", "code": "A05", "activity": "ITB Study", "unit": "Lot", "qty": 1.0, "baseIntUnit": 30.0, "baseExtUnit": 0.0, "std": null}, {"part": "tel", "code": "A06", "activity": "Technical Clarification", "unit": "Lot", "qty": 1.0, "baseIntUnit": 30.0, "baseExtUnit": 0.0, "std": {"activity": "Technical Clarification", "unit": "Lot", "int": {"상": 0.0, "중": 30.0, "하": 15.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "tel", "code": "A09", "activity": "Telecom.&Security Specification", "unit": "Ea", "qty": 7.0, "baseIntUnit": 15.0, "baseExtUnit": 0.0, "std": {"activity": "Telecom.&Security Specification", "unit": "Set", "int": {"상": 0.0, "중": 15.0, "하": 10.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n - 중 : New Generation\\n - 하 : HEC Standard Specification 적용."}}, {"part": "tel", "code": "A07", "activity": "Telecom.&Security System", "unit": "Set", "qty": 2.0, "baseIntUnit": 30.0, "baseExtUnit": 0.0, "std": {"activity": "Telecom.&Security System", "unit": "Set", "int": {"상": 0.0, "중": 30.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "tel", "code": "A08", "activity": "Telecom.&Security Bulk Material", "unit": "Set", "qty": 1.0, "baseIntUnit": 10.0, "baseExtUnit": 0.0, "std": {"activity": "Telecom.&Security Bulk Material", "unit": "Set", "int": {"상": 0.0, "중": 10.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}, {"part": "tel", "code": "A10", "activity": "Telecom.&Security Block Diagram", "unit": "Ea", "qty": 7.0, "baseIntUnit": 15.0, "baseExtUnit": 5.0, "std": {"activity": "Telecom.&Security Block Diagram", "unit": "Set", "int": {"상": 0.0, "중": 15.0, "하": 10.0}, "ext": {"상": 0.0, "중": 5.0, "하": 5.0}, "guide": "1) Draft only, design은 HEC Engineer가 수행\\n2) 난이도 산정기준\\n - 중 : ITB에 해당 도면이 없어 Concept 확정 필요\\n - 하 : ITB에 해당 도면이 포함되어 있으나 상당부분 Upgrade가 필요"}}, {"part": "tel", "code": "A11", "activity": "Telecom.&Security Cable Routing layout", "unit": "A3", "qty": 11.166666666666668, "baseIntUnit": 4.0, "baseExtUnit": 20.0, "std": {"activity": "Telecom.&Security Cable Routing layout", "unit": "A1", "int": {"상": 0.0, "중": 4.0, "하": 0.0}, "ext": {"상": 0.0, "중": 20.0, "하": 0.0}, "guide": "1) JB, LCP, Equipment 위치 표현\\n2) Cable Tray Route 및 Typical Tray Section 표현\\n3) Cable Trench, Ductbank Route 및 Typical Trench, Ductbank Section 표현\\n4) 난이도는 일괄 \\"중\\" 적용"}}, {"part": "tel", "code": "A12", "activity": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "unit": "A3", "qty": 5.583333333333334, "baseIntUnit": 3.0, "baseExtUnit": 10.0, "std": {"activity": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "unit": "A3", "int": {"상": 3.0, "중": 2.5, "하": 0.0}, "ext": {"상": 15.0, "중": 10.0, "하": 0.0}, "guide": "1) Tray Size, Routing, Layer, Main Support 표현\\n2) Trench Detail, Size, Routing 표현 \\n3) Duct Bank Size, Location 표현\\n4) 난이도는 \\"중\\" 적용 (단, 사우디, 쿠웨이트, UAE 사업은 \\"상\\"적용)"}}, {"part": "tel", "code": "A13", "activity": "Telecom.&Security Cable Schedule", "unit": "A3", "qty": 18.75, "baseIntUnit": 0.05, "baseExtUnit": 0.1, "std": {"activity": "Telecom.&Security Cable Schedule", "unit": "Ea", "int": {"상": 0.0, "중": 0.05, "하": 0.0}, "ext": {"상": 0.0, "중": 0.1, "하": 0.0}, "guide": "1) Cable no, From/To, Cable code, \\n2) Cable gland spec. & q\'ty\\n3) 난이도는 일괄 \\"중\\" 적용"}}, {"part": "tel", "code": "A14", "activity": "Telecom.&Security Typical Installation Details", "unit": "A3", "qty": 7.0, "baseIntUnit": 0.2, "baseExtUnit": 2.0, "std": {"activity": "Telecom.&Security Typical Installation Details", "unit": "A3", "int": {"상": 0.0, "중": 0.2, "하": 0.0}, "ext": {"상": 0.0, "중": 2.0, "하": 0.0}, "guide": "1) 통신 장비별 Type Assign Index (Excel) 작성 필요\\n2) 난이도는 일괄 \\"중\\" 적용"}}, {"part": "tel", "code": "A15", "activity": "BM Take-off & BOQ작성", "unit": "Lot", "qty": 1.0, "baseIntUnit": 80.0, "baseExtUnit": 200.0, "std": {"activity": "BM TAKE-OFF & BOQ작성", "unit": "Lot", "int": {"상": 0.0, "중": 0.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 최대 3차 MTO까지 산출\\n2) Cable은 Cable Schedule 작성 후 산출\\n3) Cable Tray & Fitting은 실 도면 및 Cable Route Sketch기준으로 산출\\n4) Cable Gland는 Cable schedule에서 산출\\n5) 통신 장비별 자재는 BM Sketch 기준으로 산출"}}, {"part": "tel", "code": "A16", "activity": "타부서 Inform 용 자료작성", "unit": "Lot", "qty": 1.0, "baseIntUnit": 50.0, "baseExtUnit": 100.0, "std": {"activity": "타부서 INFORM용 자료작성", "unit": "Lot", "int": {"상": 0.0, "중": 50.0, "하": 0.0}, "ext": {"상": 0.0, "중": 100.0, "하": 0.0}, "guide": "1) Cable Tray & Duct Loading Data\\n2) Cable Trench, Ductbank & Foundation Data\\n3) Electrical Load List\\n4) Heat Dissipation Data\\n5) Field Grounding Inform\\n6) 난이도는 일괄 \\"중\\" 적용"}}, {"part": "tel", "code": "A17", "activity": "Cold Eye Review", "unit": "Lot", "qty": 1.0, "baseIntUnit": 10.0, "baseExtUnit": 0.0, "std": {"activity": "COLD EYE REVIEW", "unit": "Lot", "int": {"상": 0.0, "중": 10.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1일 소요 기준"}}, {"part": "tel", "code": "A18", "activity": "3D Modeling", "unit": "Lot", "qty": 1.0, "baseIntUnit": 10.0, "baseExtUnit": 60.0, "std": {"activity": "3D Modeling", "unit": "Lot", "int": {"상": 0.0, "중": 10.0, "하": 0.0}, "ext": {"상": 0.0, "중": 60.0, "하": 0.0}, "guide": ""}}, {"part": "tel", "code": "A19", "activity": "Study & Calculation", "unit": "Set", "qty": 1.0, "baseIntUnit": 20.0, "baseExtUnit": 0.0, "std": {"activity": "Study & Calculation", "unit": "Set", "int": {"상": 0.0, "중": 20.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 건별로 내부MH 할당하고, 필요 시 3rd Party에 의뢰\\n2) Study & Caculation 항목\\n - Wireless Coverage Study\\n - Radio Coverage Study\\n - CCTV Coverage Study\\n - Sound Coverage Study & Review\\n - Link Budget Calculation"}}, {"part": "tel", "code": "A20", "activity": "Telecom. Tie-in List", "unit": "Lot", "qty": 1.0, "baseIntUnit": 20.0, "baseExtUnit": 0.0, "std": {"activity": "Telecom. Tie-in List", "unit": "Lot", "int": {"상": 0.0, "중": 20.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}], "std": {"ci": {"meetingcoordination": {"activity": "Meeting & Coordination", "unit": "Lot", "int": {"SPI": 0.0, "상": 220.0, "중": 150.0, "하": 80.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정 기준\\n- Project 표준 난이도를 적용"}, "reviewinterfacework": {"activity": "Review & Interface work.", "unit": "Lot", "int": {"SPI": 0.0, "상": 150.0, "중": 100.0, "하": 50.0}, "ext": {"SPI": 0.0, "상": 250.0, "중": 200.0, "하": 100.0}, "guide": ""}, "발주처협의technicalclarification": {"activity": "발주처 협의 (Technical Clarification)", "unit": "Lot", "int": {"SPI": 0.0, "상": 150.0, "중": 100.0, "하": 50.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "instrumentationautomationdesigncriteria": {"activity": "Instrumentation & Automation Design Criteria", "unit": "Set", "int": {"SPI": 0.0, "상": 200.0, "중": 160.0, "하": 80.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n - 상 : New Generation 및 Project 표준 난이도 \'상\'.\\n - 중 : New Generation 및 Project 표준 난이도 \'중\' / \'하\'.\\n - 하 : HEC Standard Specification 적용."}, "controlsystem": {"activity": "Control System", "unit": "Set", "int": {"SPI": 0.0, "상": 36.0, "중": 30.0, "하": 24.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n- Project 표준 난이도를 적용"}, "packagedinstrumentmeteringanalzyertgsetc": {"activity": "Packaged Instrument (Metering, Analzyer, TGS & etc.)", "unit": "Set", "int": {"SPI": 0.0, "상": 130.0, "중": 100.0, "하": 40.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "fieldinstrument": {"activity": "Field Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 12.0, "중": 10.0, "하": 8.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "others": {"activity": "Others", "unit": "Set", "int": {"SPI": 0.0, "상": 80.0, "중": 60.0, "하": 20.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "cvvalvecoefficientcalculationforsizing": {"activity": "Cv (Valve Coefficient) Calculation for Sizing", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 0.4, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 별도 성과품을 제출하거나, 직접 수행하는 경우만 M/H 반영.\\n2) 난이도는 일괄 \\"중\\" 적용."}, "psvorificicecalculationforsizing": {"activity": "PSV Orificice Calculation for Sizing", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 0.4, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 별도 성과품을 제출하거나, 직접 수행하는 경우만 M/H 반영.\\n2) 난이도는 일괄 \\"중\\" 적용."}, "flowelementcalculation": {"activity": "Flow Element Calculation", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 0.4, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 별도 성과품을 제출하거나, 직접 수행하는 경우만 M/H 반영.\\n2) 난이도는 일괄 \\"중\\" 적용."}, "datasheetvalveinlineinstruments": {"activity": "Data Sheet (Valve & In-line Instruments)", "unit": "Ea", "int": {"SPI": 0.5, "상": 0.3, "중": 0.2, "하": 0.2}, "ext": {"SPI": 0.3, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n- Project 표준 난이도를 적용. 단, SPI입력 시 \'SPI\'"}, "summarylistotherinstruments": {"activity": "Summary List (Other Instruments)", "unit": "Ea", "int": {"SPI": 0.0, "상": 0.06, "중": 0.05, "하": 0.05}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "packagedinstrument": {"activity": "Packaged Instrument", "unit": "Set", "int": {"SPI": 0.0, "상": 24.0, "중": 20.0, "하": 16.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "valves": {"activity": "Valves", "unit": "Set", "int": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "controlsystemconfiguration": {"activity": "Control System Configuration", "unit": "A3", "int": {"SPI": 0.0, "상": 50.0, "중": 40.0, "하": 30.0}, "ext": {"SPI": 0.0, "상": 4.8, "중": 4.0, "하": 3.2}, "guide": "1) Draft only, design은 HEC Engineer가 수행\\n2) 난이도 산정기준\\n - 상 : ITB에 해당 도면이 없어 Concept 확정 필요 및 Project 표준 난이도 \'상\'\\n - 중 : ITB에 해당 도면이 없어 Concept 확정 필요 및 Project 표준 난이도 \'중\' / \'하\'\\n - 하 : ITB에 해당 도면이 포함되어 있으나 상당부분 Upgrade가 필요"}, "instrumentcableroutinglayout": {"activity": "Instrument Cable Routing layout", "unit": "A1", "int": {"SPI": 0.0, "상": 5.0, "중": 4.0, "하": 3.0}, "ext": {"SPI": 0.0, "상": 42.0, "중": 35.0, "하": 28.0}, "guide": "1) JB, LCP위치 표현\\n2) Cable Tray Route 및 Typical Tray Section 표현\\n3) Cable Trench, Ductbank Route 및 Typical Trench, Ductbank Section 표현\\n4) 난이도는 Project 표준 난이도를 적용"}, "instrumentcabletraytrenchductbanklayout": {"activity": "Instrument Cable Tray, Trench, Ductbank Layout", "unit": "A3", "int": {"SPI": 0.0, "상": 3.0, "중": 2.5, "하": 2.0}, "ext": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "guide": "1) Tray Size, Routing, Layer, Main Support 표현\\n2) Trench Detail, Size, Routing 표현 \\n3) Duct Bank Size, Location 표현\\n4) 난이도는 Project 표준 난이도를 적용"}, "jbwiringlayoutsketcheddrawing": {"activity": "JB & Wiring Layout (Sketched Drawing)", "unit": "A3", "int": {"SPI": 0.0, "상": 3.0, "중": 2.5, "하": 2.0}, "ext": {"SPI": 0.0, "상": 18.0, "중": 15.0, "하": 12.0}, "guide": "1) Multi Cable & Tray Routing 표현\\n2) 계기, JB (PKG JB포함), LCP, LCS 위치 표현 \\n3) 난이도는 Project 표준 난이도를 적용"}, "controlroompibequipmentarrangementlayout": {"activity": "Control Room / PIB Equipment Arrangement layout", "unit": "A3", "int": {"SPI": 0.0, "상": 5.0, "중": 4.0, "하": 3.0}, "ext": {"SPI": 0.0, "상": 36.0, "중": 30.0, "하": 24.0}, "guide": "1) Cabinet Type, Size, Scope (Package, Electric 등) 명기\\n2) 난이도 산정기준\\n - 상 : ITB에 해당 도면이 없어 도면 작성 필요 및 Project 표준 난이도 \'상\'\\n - 중 : ITB에 해당 도면이 없어 도면 작성 필요 및 Project 표준 난이도 \'중\' / \'하\'\\n - 하 : ITB에 해당 도면이 포함되어 있을 경우"}, "instrumentindexandiolist": {"activity": "Instrument Index and I/O List", "unit": "Ea", "int": {"SPI": 0.5, "상": 0.36, "중": 0.3, "하": 0.24}, "ext": {"SPI": 0.3, "상": 0.18, "중": 0.15, "하": 0.12}, "guide": "1) Tag no., Inst. Type, P&ID no. \\n2) Piping Class, Line Size,  Location, System, I/O Type, etc.\\n3) 난이도는 Project 표준 난이도를 적용. 단, SPI입력 시 \'SPI\'"}, "instrumentcableschedule": {"activity": "Instrument Cable Schedule", "unit": "Ea", "int": {"SPI": 0.0, "상": 0.06, "중": 0.05, "하": 0.04}, "ext": {"SPI": 0.0, "상": 0.12, "중": 0.1, "하": 0.08}, "guide": "1) Cable no, From/To, Cable code, \\n2) Cable gland spec. & q\'ty\\n3) JB & Wiring Layout에 의한 Multi cable length적용\\n4) 난이도는 Project 표준 난이도를 적용"}, "typicalhookupdrawings": {"activity": "Typical Hook-up drawings", "unit": "A3", "int": {"SPI": 0.0, "상": 0.36, "중": 0.3, "하": 0.24}, "ext": {"SPI": 0.0, "상": 3.6, "중": 3.0, "하": 2.6}, "guide": "1) Tag 별 Type Assign Index (Excel) 작성 필요\\n2) 난이도는 Project 표준 난이도를 적용"}, "instrumenttypicalinstallationdetails": {"activity": "Instrument Typical Installation Details", "unit": "A3", "int": {"SPI": 0.0, "상": 0.24, "중": 0.2, "하": 0.16}, "ext": {"SPI": 0.0, "상": 2.4, "중": 2.0, "하": 1.6}, "guide": "1) Tag 별 Type Assign Index (Excel) 작성 필요\\n2) 난이도는 Project 표준 난이도를 적용"}, "3rdpartysysteminterfacedesigncableblockdiagram": {"activity": "3\'rd Party System Interface Design & Cable Block Diagram", "unit": "", "int": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "major3rdpartysysteminterface": {"activity": "- Major 3rd Party System Interface", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 5.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 5.0, "하": 0.0}, "guide": "1) Skid JB , LCP, LCS, Loosed Instrument informs\\n2) HEC Scope Cable에 대한 Cable informs\\n3) Vendor Data 및 Reference Project 자료 검토"}, "minor3rdpartysysteminterface": {"activity": "- Minor 3rd Party System Interface", "unit": "Set", "int": {"SPI": 0.0, "상": 0.0, "중": 2.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 2.0, "하": 0.0}, "guide": ""}, "bmtakeoffboq작성": {"activity": "BM Take-off & BOQ작성", "unit": "Lot", "int": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 최대 3차 MTO까지 산출\\n2) Cable은 Cable Schedule 작성 후 산출\\n3) Ladder Tray & Fitting은 실 도면을 기준으로 산출\\n4) Perporate Tray는 Cable Route Sketch 작성 후 산출\\n5) Cable Gland는 Cable schedule에서 산출\\n6) JB, Air Pipe & Fittings는 BM Sketch 작성 후 산출\\n7) Tubing, Hook-up Valve, Fitting, Gasket & Bolt는 Hook-up물량 산출표 (Excel)를 작성하여 산출"}, "타부서inform용자료작성": {"activity": "타부서 Inform용 자료작성", "unit": "Lot", "int": {"SPI": 0.0, "상": 120.0, "중": 100.0, "하": 80.0}, "ext": {"SPI": 0.0, "상": 200.0, "중": 190.0, "하": 180.0}, "guide": "1) Cable Tray & Duct Loading Data\\n2) Cable Trench, Ductbank & Foundation Data\\n3) Electrical Load List\\n4) Heat Dissipation Data\\n5) Heat Tracing Instrument List (ET, ST, CT)\\n6) Field Grounding Inform\\n7) 난이도는 Project 표준 난이도를 적용"}, "3dmodeling": {"activity": "3D Modeling", "unit": "Lot", "int": {"SPI": 0.0, "상": 40.0, "중": 30.0, "하": 20.0}, "ext": {"SPI": 0.0, "상": 300.0, "중": 250.0, "하": 200.0}, "guide": "1) 난이도는 Modeling & Review 역무 Volume에 따라 난이도를 적용"}, "spiseeddb": {"activity": "SPI SEED DB", "unit": "Lot", "int": {"SPI": 0.0, "상": 100.0, "중": 80.0, "하": 60.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도는 Project 표준 난이도를 적용"}, "coldeyereview": {"activity": "Cold Eye Review", "unit": "Lot", "int": {"SPI": 0.0, "상": 30.0, "중": 20.0, "하": 15.0}, "ext": {"SPI": 0.0, "상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도는 Project 표준 난이도를 적용"}}, "tel": {"meetingcoordination": {"activity": "Meeting & Coordination", "unit": "Lot", "int": {"상": 0.0, "중": 60.0, "하": 30.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정 기준\\n - 중 : ITB가 불분명하여 확인할 사항이 많은 경우\\n - 하 : ITB 자료가 충분한 경우"}, "reviewinterfacework": {"activity": "Review & Interface work.", "unit": "Lot", "int": {"상": 0.0, "중": 30.0, "하": 15.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "technicalclarification": {"activity": "Technical Clarification", "unit": "Lot", "int": {"상": 0.0, "중": 30.0, "하": 15.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "telecomsecurityspecification": {"activity": "Telecom.&Security Specification", "unit": "Set", "int": {"상": 0.0, "중": 15.0, "하": 10.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 난이도 산정기준\\n - 중 : New Generation\\n - 하 : HEC Standard Specification 적용."}, "telecomsecuritysystem": {"activity": "Telecom.&Security System", "unit": "Set", "int": {"상": 0.0, "중": 30.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "telecomsecuritybulkmaterial": {"activity": "Telecom.&Security Bulk Material", "unit": "Set", "int": {"상": 0.0, "중": 10.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}, "telecomsecurityblockdiagram": {"activity": "Telecom.&Security Block Diagram", "unit": "Set", "int": {"상": 0.0, "중": 15.0, "하": 10.0}, "ext": {"상": 0.0, "중": 5.0, "하": 5.0}, "guide": "1) Draft only, design은 HEC Engineer가 수행\\n2) 난이도 산정기준\\n - 중 : ITB에 해당 도면이 없어 Concept 확정 필요\\n - 하 : ITB에 해당 도면이 포함되어 있으나 상당부분 Upgrade가 필요"}, "telecomsecuritycableroutinglayout": {"activity": "Telecom.&Security Cable Routing layout", "unit": "A1", "int": {"상": 0.0, "중": 4.0, "하": 0.0}, "ext": {"상": 0.0, "중": 20.0, "하": 0.0}, "guide": "1) JB, LCP, Equipment 위치 표현\\n2) Cable Tray Route 및 Typical Tray Section 표현\\n3) Cable Trench, Ductbank Route 및 Typical Trench, Ductbank Section 표현\\n4) 난이도는 일괄 \\"중\\" 적용"}, "telecomsecuritycabletraytrenchductbanklayout": {"activity": "Telecom.&Security Cable Tray, Trench, Ductbank Layout", "unit": "A3", "int": {"상": 3.0, "중": 2.5, "하": 0.0}, "ext": {"상": 15.0, "중": 10.0, "하": 0.0}, "guide": "1) Tray Size, Routing, Layer, Main Support 표현\\n2) Trench Detail, Size, Routing 표현 \\n3) Duct Bank Size, Location 표현\\n4) 난이도는 \\"중\\" 적용 (단, 사우디, 쿠웨이트, UAE 사업은 \\"상\\"적용)"}, "telecomsecuritycableschedule": {"activity": "Telecom.&Security Cable Schedule", "unit": "Ea", "int": {"상": 0.0, "중": 0.05, "하": 0.0}, "ext": {"상": 0.0, "중": 0.1, "하": 0.0}, "guide": "1) Cable no, From/To, Cable code, \\n2) Cable gland spec. & q\'ty\\n3) 난이도는 일괄 \\"중\\" 적용"}, "telecomsecuritytypicalinstallationdetails": {"activity": "Telecom.&Security Typical Installation Details", "unit": "A3", "int": {"상": 0.0, "중": 0.2, "하": 0.0}, "ext": {"상": 0.0, "중": 2.0, "하": 0.0}, "guide": "1) 통신 장비별 Type Assign Index (Excel) 작성 필요\\n2) 난이도는 일괄 \\"중\\" 적용"}, "bmtakeoffboq작성": {"activity": "BM TAKE-OFF & BOQ작성", "unit": "Lot", "int": {"상": 0.0, "중": 0.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 최대 3차 MTO까지 산출\\n2) Cable은 Cable Schedule 작성 후 산출\\n3) Cable Tray & Fitting은 실 도면 및 Cable Route Sketch기준으로 산출\\n4) Cable Gland는 Cable schedule에서 산출\\n5) 통신 장비별 자재는 BM Sketch 기준으로 산출"}, "타부서inform용자료작성": {"activity": "타부서 INFORM용 자료작성", "unit": "Lot", "int": {"상": 0.0, "중": 50.0, "하": 0.0}, "ext": {"상": 0.0, "중": 100.0, "하": 0.0}, "guide": "1) Cable Tray & Duct Loading Data\\n2) Cable Trench, Ductbank & Foundation Data\\n3) Electrical Load List\\n4) Heat Dissipation Data\\n5) Field Grounding Inform\\n6) 난이도는 일괄 \\"중\\" 적용"}, "coldeyereview": {"activity": "COLD EYE REVIEW", "unit": "Lot", "int": {"상": 0.0, "중": 10.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1일 소요 기준"}, "3dmodeling": {"activity": "3D Modeling", "unit": "Lot", "int": {"상": 0.0, "중": 10.0, "하": 0.0}, "ext": {"상": 0.0, "중": 60.0, "하": 0.0}, "guide": ""}, "studycalculation": {"activity": "Study & Calculation", "unit": "Set", "int": {"상": 0.0, "중": 20.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": "1) 건별로 내부MH 할당하고, 필요 시 3rd Party에 의뢰\\n2) Study & Caculation 항목\\n - Wireless Coverage Study\\n - Radio Coverage Study\\n - CCTV Coverage Study\\n - Sound Coverage Study & Review\\n - Link Budget Calculation"}, "telecomtieinlist": {"activity": "Telecom. Tie-in List", "unit": "Lot", "int": {"상": 0.0, "중": 20.0, "하": 0.0}, "ext": {"상": 0.0, "중": 0.0, "하": 0.0}, "guide": ""}}}}')
HYUNDAI_LOGO_B64='iVBORw0KGgoAAAANSUhEUgAAAJcAAAAeCAYAAAAhIIxpAAAAAXNSR0IArs4c6QAAAAlwSFlzAAAOxAAADsQBlSsOGwAAABl0RVh0U29mdHdhcmUATWljcm9zb2Z0IE9mZmljZX/tNXEAABX5SURBVHhe7VwJfJXVlT/f8pbkJeRFsAOCqID4o2IJ1gBBo61lLG7U6qitdFrrOC7V6YhYcaH5/FJEhQriMrYuuLRMBYrSqBQd7DhDNYSo9dda0kGgFZSwKAm89/KWb7nzP/d772V7W1imvzpef+El37v33HPPOffsn7ppmvTZ+IwCR4IC+sECvWNq0wSy3SHzN170GsMwRrd+m/xly8y28c7Bwvxs3aeLAgctXIptL1I0ZQzIcZwx7t3TKRx6hvbFP8bfaw6VREZrYDgJ331EKpFeRpSMLjPrYr/OBddoDvtJse8nzVdNmp/I6vpvWjPiCTp/xxLSAtXkJAGjHDBiSwCjNR9uRnOongKhaynVReQPKGSn2unlWT+g8+7/PvlCteRYRI69i9bM+oHZ2Chy4tISuph8FReTHSdyHZs0+xazNsk0IaMldCfOMo6cRAnkUQQpvp1kx1abdcnmnguMlor5pAdGynMVG4rqkuLfTlbHKrPO+V3f6cDpq6DZdHKto4uBkt/rIdA3+ro5OfZEZr6xsaIRdB6F84I+Toys1GyzPhmV00sC2mfSHVNWT9Mq9LNJI5p7xsorjT1jryAdgkDKPKPB+LXZaOYkfsl72fpRFAzNJAUwfVK42rA2p3CRvdNH/s9dRf5QOelSuCpMs+0xo6VyOwXDN1IK59QD/Hy8MW7cqWZbm9uPyCygmr2U9OAYUhTsCWG09lxvmo0CDJgB2GeTaxPFO/fSypW3Yn3u8ynKJApUzCQVeNtgvrDZ55DCBdpcAjgTyQbRShl8dlW91WjWv4FLsTy7RKGZgDOS7BJZp/B+VXOMVut8szbyalYomitnU1nox8ALeIJupQw/hCsV5Y2zwgVKXA58xkrhShxgKHfi5+CEq2GcoQl14iKJS8yiJ7vCT7GCob248VX+ifSLr38Lf/2sFFwLzHEoFcO5AVhAFlSCKsg3qpnRHZhfjhuINUIezJwcWWi0uNPJX3E2dX1CFKicQEu330B19FA/SKo1m/yDxlAc8/yVIM2eJnNK7CdpoYhIXFi4iDoLnkuIuJxrgRbCSQHvnoJ8QH7HWo0Zzpcm37Ch3aD4wDTMVX5oGA0rWNDldCE6AQfChTkq+KwH88OxMIfxDlTqlEox06VwGesDleQXd0I7e5eANT5fzFxDYFuLyZ++T4oi6dtj7Pfow96Q6ICEYSIEFqNE8e8BqnridUqVdgp1WtQhNLorVY0v0zRM4VMXPzKGfrTS3DW8FP1fkFeH/KVqXwcC/h5MCJLF5lFvNFrFCrNW2Z2BbbSGR5Iq5kqBUH3M/ChZ9o2HvHc+ACwQrrOXrNjq3MyESAo6lzT9GIJ8kqoMpWmPQeJJqoXs8OB8RMmul4G/x83sgFkVLg6jXIizD0mb0JHGe+M0c3ybA64fg6nVUvA0THOtzWRZ/5UHHx9k5XIgUuA25D7sgISr4dSmIUIVjUoXhEhxqbHrKNpp8RnSPnwCyFYFjiPa9S+0a/jCI8agEgHD33kfZsWkskH3UBK8CVSGKWHDl4tdmQXhpBZRYFA5pSKe1krtvw0+w44Stxj4NBZg4W4xJ0WuybcYfsxD8M9ulNpJuFDf7Uzk3tM1mHon9ntzSvTavHBaKpbBJbhCanTW/5FKtpEO2RCvjKJizZeKPGpOij1QAJ/psCJHVrgcn/0jbZDvKOpI0SYnSA8nB3ULVgazLgiYot5unND6lPnn2rS/MXAe9LiF/Gte7xWC0GW06Pl9vKtGLqCnP7wcflSNVN++4HeMZnocfswbRnPZOTA9l3jPpT/xhjk5/sihYFt8LVAVooRLnTkStBAF8dOXBNL6QOgKDIVtWnZ0/673pJd8nNcR9EwoBLOPbBc/5wDMIqceVL96DR1g4SG6OV5NtsClggbrNSxosSpEaba4A89vLgWJgnOE1IqjEc3V4L6xd9djABcbV09VcA9zyxc78IYbugYmpIUdGOnHqcoiEG0a+bX75Tp+5qRsmJe8WuCQzzEQAIgVBzI971zB0UlJo9R5JQHLTCrhBnlTNddZrPgREqaS9GKqgl5JsbPZL/DyJsPRp4D6PeOU5ofNP9RtGxBGfSezM6los0hTZvWHww4tnrIzKR1umIocg1MQRgs9APM3i5Iwf4o6ify+1wB3vHSwA9DAiQPz4cT/8ZBw/ZtcLOUKDmfuAcsQMTb68jC68IFLEq65U1b/g1Klf5kiNqXgxN8CrVVwWMClKhCgDuduzPvmYaF5yZcwz24pey4p0Yvgg5yQjpCQu4LD7OXRNtGJvnmHBc+/NSAcLZLyj8aGivHQl30sA+fbZJ6iSkbtAxxFhath6jNBVwxaSCww8AUfSFTTZhtOYMaJz7UhX4YoGFemf8M4sfV+8/3atwaIV+/pAmpJOHlOx5sp7CUXVO2eb6Z+D2oK+TJMZcFigWWiue515uCI9Hr/3w0XdNACUxA1Tsl59mwq4ggIl+tW3Koe5TuenfiPIMTzEuH85rAndi58GT/8xISK6Iy+ctBMkwnN2L1IEzyeE4YdL8N3ryPkHlpsDzjra40N6jLyl8+UuRvWWlb0p2Zd1/piaz/13/fy/Q/PaQtqroapTSOFps2hKDvxgu6IhSniYkkhrdUTrxhuRaX/bOOkd881/6cmd4a92Dk8c7jHrO3cnmuqYRgKnfdjL8ooZSjqG/C5ZnZPVX5byrJP7RxOjbipd6DJN4GEvYnICkJRWEa+hs8C2drc1CkoXI5jL9CqoDo6U7TBLqNnk8gDlSpYvB8HcDKIc+9D6eWVXKWXEpmWP+Sedi/Uz4Acsr5EGgjRnCJnQGKqRCEv8eD5p8l9iplytv2Fd5J5rtiTqGj8W76JyLu1I/gpahn6rs8rXA1nrqon3Xe5l3pQaJZ04hnRAZYNu3D+Kv8pFPnFt4lqnj5kmv61AEiHVlSh1ngxuaroVduwccVV6QyfAy2QxvAICxlHx4KGI0WDIEXtvRnjw3G0QrXpBCrjpNDYzTmYx49c+B65x2HPczXA1Liv1ixRObJHnulZJEs3WKwgBu7USZSTSBUIt9EY995ytOQUqBP+tSSnhH2ZmYoyjLTgqn6zOZ5i5clzZM2P025OnBLVqLXtKgH4QUzhrghV+wLqgi/kxUd2cnDZi5PrqU5zMGp2/4cjp+ZyX5l4tRrWJtJ+iyJIJt8utdYANVbPQ7BwVQWOpc7ETXh8T9Hz2bhlZVzxgDKQta8iNVAhgnIe/1jZwka+bfTsXG9+kYgZUYzEIW2ZvaRu/5EhD5sZfwUyR588Z9bv6lkPDEg4TEc7lTshl4GqIPrluTKSRSmEEpDcrHLy4GRSA4XwYYFnweIiuZ18Oot0X/oWTabD38rQt18LheLhw7ySuHYj2o+wcOLDQhHzKM5ayqV5cdQPHdzEgfhauYjP5lFT5xjjWp8w22r3FhSwIMr1trJDagNVY6pCA+QZ1R2CokdvQ79Vl3RIBbUXEd4OOK8fkm3BlFm8QWfh+cpOwPbmFxoyI8J5Ib2DEh1ryEVNs+dQaAf2HeFJl8gZnGSnC3ePxNFJ8VzgF4REZ8o/ygfApxz4FEFb4uPCWO+h+L7n0IO1OLtgIPSVXQliK3AZAnh8yj6qGGdxLK8ITmI/+leyt6+fcNmOc5c+2Pc56kjSZjeAvBbXDw/SHPY8PndMhANVaKWYi8f/WogyXHAGMUd6c5BRLzDM8eUws7FJRhMijrdHJU1zT2FkrxyxFE657EcymmDr54xKUhu3i+Ue5uTo1TzPnJHsZ86NZWEfTRjmyk6DIsOcFL2M25yM9UODpI8GjhsL7BlDm80tDTTtp9BwuFfrbktSfVpWXX2G+cVP+tVa4RsBbplt1nXK3qCC46U5W2jafWO5MYJCHQolyvJKKvKDKLwmTy1AnwvT7VtyimHcpdA5Hq69hKvhzKbP4zrfmKkf3tJVjYx8jvphMeRzfS8TqziDX70e2ushaK8t+cCgLQanHnwGuUmoWTirakAjN3YAeaqX4cCeBadZM0+3fiMP0xw6DR8BcwYK0aO2hY23h9TDBKA2xU/RUWpF3vK6I9jxtf9Ize0fGGVDzyUr8aY5o3OvsbRdN6jya6Rrb1IibqFJEY2B2JfvvF+ak53mjMg6rD8dnZhDgYfXT8AO6cjYSxTH+tahZ5IdQY0TuOpl6EXb/w46SLcazYEaPBuexrsefo8GU/k6X37Aq8UOFdAo/4m2n6NBGMCIhWTDjR5y6YIlG83aXbhkGBnBwq8QniTWjkKj30RKxmGP2MLGImY9aPNJmQ+wzoN5qsJDwClLglbNSOP00jbp3jA4h5nHSeATqEOukPH5jwxf0OU7mvz+MxBNcvsUO3k4o4besGATBTsqSa3Ed3FPhspC8KFi7+CdjCxfewmXk7QXaWGfRpEUaoflsoZ4WLRWBltbloV81OnMxyPc5DwjERxPlThAMrUfNEoCB4imtgmz0buEpGz53002mjsmomb4LmQPZR06Ht/VUMy3mnzKWVI9s1Mi0BLq+mcbxm1b6PzFL0Bgbsd1ehoZ6dUo+WzAmjraiyTssNDzEKgL8V0HEqwrKWHtQxEbYR9ulnB53jqY9MdJdcaRq4AjXADHfzr9AeZpFIUqV+M5VKxAJcGqgpPdjj2PpfMW/TNmIUomMEJdSGVDJxut+75o1sbeAYjb8fxE/JwCk7KcfP4vA0Yn/kZrMva1nR/id0+4egy0gA8BLm+hnRgmRURQ9FXI1bZK2my1F6Dp8Sa0FzEcBhUG7N8hDXRa0TSQ4gMdBeMDjZYeujMF+D8MGpehq0RD4jkJ4YdmjGxAxIyu2/JlCGA8nBkfvxKBBh0DbbePIWSFq2HKqgtpkO+rXD+0cLab4Wsd9iG1lywLXYq++0lmW01u28AtIVxQVsUUMOJPvfEQcJLhiqjKUjxndb0fR0vbTmUibtIDdKJ+qxcZQaNjGKl7gzJ6k7cPH3HHQTvyFKNFXAON95ixwQcXVwXRQA5ui06lvgACfdRT3WPLAIr2yOZHr+vF7Ba9hhKQ5ZA10hyf7DRaK7+OisHzNH0JtIodxToPN+4glQ46PYW/JoBpWJRtAET7c9fD8NNmeWatb7NnLwoMRsQKk9J1CTz1l8h9X6FtMLOTeA/8mzqwHnCmyRV27HYS+p302DY/1VPh5k2FicUXpMewteV0jLqKPhTnwK/6FcpktRRKbaGX5iTwbsE0indAS1ujafvOCI0dcQFpZc+jWTQMMncLV8OkZX6hlt/v+ewOPYISz6Zi9cODlTxOwUA5Ujy5ACC+lBuM7cigxFVXI68Uh5lQoJrvgKZaAyaFcajnkNSDcITAaIHeZGW0hAMjiH9206b40UYw/AqK1GhRPvAirTv+PjpvBzvHPAvaCNowGUPflroYJua3EGJulfAOrwI3v/819Igl0UKMfRM3YF+Uh9DC6wtciu6KKXiO+xt/Gybku3INR0ox/StoTPwYsCaQhn22bRU05jgG6wUCCoq/ic6V+OU09JHdCC0BvDORFWs8ON46DTY2hn+DfaIQnleR2GTt1XdAIzOj1CgEETd1MPRvZ3qOhJPynvNLIWUc5Bx0+iHtv9nGzpB3S3U7igsUp/GN/MIJHqBmVE1xs36wZbTqMQ+v7tZuqblcPXSTGtZP5PrhbjDVRHH6sJrDvuThstAg/1koC12AstBLOQjoNacpYgWIBQ2Cbkzd957HJPRuKRwxpR6AhkDbjIo2GTdNXXYQBUHr7DRaQwvhL80loX0JL2zcg7dmurdR1HJSkktw8BGkac9BowiQhBMeEAUZDzyNHzQ6ohhu+zan94Uf5mAv8SxzFnt6ER+v83Gqwn5SJlKD4UoI0XJzJgi+AQmvTDAkgLfKkZVYCN/mdRwEsBzWXixiUgThG+2Gz3Qf8L4N+5yNb/oLl22jRRDEsd3++ULuanC7ozWs57dEQLudgA4hLDS490vxXhToN5gyXr9EtxvlwhzD65DCxoPx0Xv3HOoNU5cfJdSAQV2QOty4O6Nh6uRGSUmUI5xlVtwFxqTmtebGut6HsoFlEJrLSW0Hsd6HOgaTreOBEBgKJgkxDA5zC/qMllPlMVdR5MP062xMIO0caKM/g6670YfOGZ1M9AjvnAuj7CvBCY45qGXZ11BA20WhISG8nGHDm9AILiFejsC+7k4ZTOhJ1oq7sSeuqvsRoG2GFoWzH/IZ66vgZ8Q0vsDkilPJtjtJPfCml1FlEUS4DSfAY5YIQFaBdxS9ZRX/TlXDr6b9H671vuOgSfl74A1Btnaj178n3r15rQddnAt+vAYtGhqBwAIClYxAizYBDk6gnYw3n65ghKCFp+EDtDvZZ7zXgTZnBBtXDbNz+l/ctaroJ2DtP2EdYFqdtOaWX8JBBxxXky9xJFM4VyZQBc56mQ5NfjUQjGHfOolXIq4gIq5A4BLVXRG4Wx2EF/v2p2gjqgDPINDQfazJj5xgOTi34HbocGAcUhPcS967rqXDT7GSfwGx5gMNvjPQCQprrjO5/xy4ebmsZMUsUveeBuayiYFQuIvxzuG1UM+PygNovg7cqJ+kucO+G3wBNPrbQcAmRG7JqNGsfhfChH5/DQ1zqQQ00AcQhCUgNLQJbr1KEBa6AMx/C/CmkuU8L1/kcB1ohDgiV3cXcnJ/oUp7n/S5mvWZIGCTsT40hnRlK9Z7mo+UbYAHgcVIWTfjDaPaLN4CeOvl15OIe3irPrzhY9/dW6rSf9m+HeRGV4DZ38R7gt+CdseFwj5ETTj/Izj/ApwBJh85N9WP88RNXIAzKOafhIBkDC3d9ijCmOZ+sBXxPgKdqeR2oTsXl0MoW2jFilVI08D/VeE7gh9liUxtizmyR1oQUtD1AhxUXAgrwVr9Y/JHf47Pi3RXUZ9Uo2CAsGhZtFqvpJSOHuWc5zpcD3eqeipho7vRtcA8t38X5NrZ79L0B0/Ktnbb0FzbdJcmAwPhu5gqh0H622D+dh1AV0QNfX+JVNfwjRqNZh8i0c9DIDcr1HSDhZvnaa4t+gSad6xDl10m4GyfROtmW1TfiDXxlcaylb+izVfh0JfiuwcRLaW1v+VgX+DJ+zoaMxNwWbMxQBB07WwQe6VC09tPgmBJwgOHtxHCn0DVeEFx5ajFNLb9QQ9vrQfe3A6brJF5MnRRQevcDbzBpBx49yE6pyLw6HKj+WSoEsituwU1Q86byb1/ifxYOpgYK+hPm4U5M2bJF4ft+Hr4knVgLUfd/Uf5yAaKtCNBhbOzX7xNF1kNt3b2esmPF+ZaVJPOt1058gVa2v5illaui/05WuNX1+g7/KnPf3NGtpGPrfJNh0uCDgFOOg/TfUsYVp0H0HMyO7PQpdrm99zSw/s+HYTWdSfJzZmY04Z18v+NAYe3rrEbxsw67MWJVPmswL59DuXBYDXfa03GoabxMjmbZnxvvOVZGKcieOcjo1m3Mb0nc60XPbAfCyBo0E0znouf5Lq88LxEcHcyOL1W4um9Mwma9aAnv1xc14dWmbPUJ6UvWaSulg+Vz55/RoHiFPhfhz8x2BUwxhkAAAAASUVORK5CYII='
P={'ci':'CI (계장)','tel':'TEL (통신)','both':'전체'}
REV={'전체':'both','CI (계장)':'ci','TEL (통신)':'tel'}
def num(v):
    try:
        if v in (None,'','-'): return 0.0
        return float(str(v).replace(',',''))
    except: return 0.0
def fmt(v):
    try: return f'{float(v):,.0f}'
    except: return str(v)
def f1(v):
    try: return f'{float(v):,.1f}'
    except: return str(v)
PRIMARY='#102A43'; ACCENT='#0F6CBD'; BG='#EAF1F8'; CARD='#FFFFFF'; HEAD='#163B5C'; TEXT='#172033'; YELLOW='#FFF2CC'; GREEN='#E2F0D9'; PURPLE='#EFE7FF'; LOCKED='#F2F2F2'; LINE='#D7E3EE'
class Model:
    def __init__(self):
        self.inputs=[dict(x) for x in DATA['inputs']]; self.outputs=[dict(x) for x in DATA['outputs']]; self.std=DATA['std']
        self.part='both'; self.case='short'; self.ratio_mode='original'; self.external_min='No'; self.project=''; self.base_mh=161; self.months=4; self.rate_short=32000; self.rate_comp=39000; self.internal_pct=0; self.external_pct=0; self.sync_common(); self.recalc()
    def find(self,p,c): return next((x for x in self.inputs if x['part']==p and x['code']==c),None)
    def val(self,p,c):
        x=self.find(p,c); return x.get('value','') if x else ''
    def setv(self,p,c,v):
        x=self.find(p,c)
        if x: x['value']=v
    def nval(self,p,c): return num(self.val(p,c))
    def sync_common(self):
        for p in ['ci','tel']: self.setv(p,'CTRL_PROJECT',self.project); self.setv(p,'CTRL_EXTERNAL_MIN',self.external_min)
        for p,c,v in [('ci','C35_VALUE',self.base_mh),('tel','C09_VALUE',self.base_mh),('ci','C36_VALUE',self.months),('tel','C10_VALUE',self.months),('ci','C37_VALUE',self.rate_short),('tel','C11_VALUE',self.rate_short),('ci','C38_VALUE',self.rate_comp),('tel','C12_VALUE',self.rate_comp)]: self.setv(p,c,v)
    def input_changed(self,x,v):
        x['value']=v; c=x['code']
        if c=='CTRL_PROJECT': self.project=v
        elif c=='CTRL_EXTERNAL_MIN':
            self.external_min=v
            if v=='Yes': self.ratio_mode='custom'; self.internal_pct=75; self.external_pct=25
            else: self.ratio_mode='original'
        elif c in ['C35_VALUE','C09_VALUE']: self.base_mh=num(v)
        elif c in ['C36_VALUE','C10_VALUE']: self.months=num(v)
        elif c in ['C37_VALUE','C11_VALUE']: self.rate_short=num(v)
        elif c in ['C38_VALUE','C12_VALUE']: self.rate_comp=num(v)
        self.sync_common(); self.recalc()
    def compute_diffs(self):
        gp=self.val('ci','C01_SEL'); gd='상' if gp=='1)' else '중' if gp=='2)' else '하' if gp=='3)' else '중'; self.setv('ci','C01_DIFF',gd)
        for c in ['C04','C06','C08','C10','C12']:
            s=self.val('ci',c+'_SEL'); self.setv('ci',c+'_DIFF','상' if s=='1)' and gd=='상' else '중' if s=='1)' else '하' if s=='2)' else '중')
        for code in ['C14','C23']:
            s=self.val('ci',code+'_SEL'); self.setv('ci',code+'_DIFF','SPI-내부' if s=='1)' else 'SPI-외부' if s=='2)' else gd)
        for c in ['C17','C18','C26','C27','C28','C32']: self.setv('ci',c+'_DIFF',gd)
        for c in ['C19','C21']:
            s=self.val('ci',c+'_SEL'); self.setv('ci',c+'_DIFF','상' if s=='1)' and gd=='상' else '중' if s=='1)' else '하' if s=='2)' else '중')
        m=self.val('ci','C29_SEL'); self.setv('ci','C29_DIFF','상' if m=='1)' else '중' if m=='2)' else '하' if m=='3)' else '중')
        self.setv('tel','C01_DIFF','중' if self.val('tel','C01_SEL')=='1)' else '하' if self.val('tel','C01_SEL')=='2)' else '중')
        self.setv('tel','C03_DIFF','중' if self.val('tel','C03_SEL')=='1)' else '하' if self.val('tel','C03_SEL')=='2)' else '중')
        self.setv('tel','C05_DIFF','상' if self.val('tel','C05_SEL')=='1)' else '중')
    def update_qty(self):
        ci1=self.nval('ci','A01_QTY'); ci2=self.nval('ci','A02_QTY'); ci3=self.nval('ci','A03_QTY'); inline=ci3 if ci3>0 else ci2*.25
        self.setv('ci','A15_QTY',round(inline)); self.setv('ci','A16_QTY',max(0,round(ci2-inline)))
        for c in ['A22','A23','A24']: self.setv('ci',c+'_QTY',math.ceil(ci1/500) if ci1 else 0)
        self.setv('ci','A26_QTY',ci2); self.setv('ci','A27_QTY',round(ci2/5)); self.setv('ci','A28_QTY',min(100,math.ceil(ci2/80) if ci2 else 0))
        sys=self.nval('tel','A01_QTY'); site=self.nval('tel','A02_QTY'); bld=self.nval('tel','A03_QTY')
        self.setv('tel','A09_QTY',sys); self.setv('tel','A10_QTY',sys); self.setv('tel','A11_QTY',round(site/90000*2+bld,2)); self.setv('tel','A12_QTY',round(site/90000+bld/2,2)); self.setv('tel','A13_QTY',round(site/10000,2)); self.setv('tel','A14_QTY',sys)
        for o in self.outputs: o['qty']=self.nval(o['part'],o['code']+'_QTY')
    def diff(self,o):
        p=o['part']; c=o['code']
        if p=='ci':
            if c in ['A04','A05','A06']: return self.val('ci','C01_DIFF')
            mp={'A07':'C04','A08':'C06','A09':'C08','A10':'C10','A11':'C12','A15':'C14','A16':'C17','A21':'C19','A25':'C21','A26':'C23','A34':'C29'}
            if c in mp: return self.val('ci',mp[c]+'_DIFF')
            if c in ['A17','A18','A19','A20']: return self.val('ci','C18_DIFF')
            if c in ['A22','A23','A24','A27','A28','A29','A30','A31']: return self.val('ci','C26_DIFF')
            if c in ['A32','A33']: return self.val('ci','C27_DIFF')
            if c in ['A35','A36']: return self.val('ci','C32_DIFF')
            return self.val('ci','C01_DIFF')
        if c in ['A04','A05','A06']: return self.val('tel','C01_DIFF')
        if c in ['A09','A10']: return self.val('tel','C03_DIFF')
        if c=='A12': return self.val('tel','C05_DIFF')
        return '중'
    def units(self,o):
        d=self.diff(o) or '중'; st=o.get('std')
        if st:
            lk='SPI' if 'SPI' in d else d; iu=st['int'].get(lk,0); eu=st['ext'].get(lk,0)
            if iu or eu: return iu,eu,d
        return o.get('baseIntUnit',0),o.get('baseExtUnit',0),d
    def visible_outputs(self): return self.outputs if self.part=='both' else [o for o in self.outputs if o['part']==self.part]
    def raw_row(self,o):
        iu,eu,d=self.units(o); q=num(o.get('qty')); return {'part':o['part'],'code':o['code'],'activity':o['activity'],'unit':o['unit'],'qty':q,'diff':d,'iu':iu,'eu':eu,'hec':iu*q,'ext':eu*q,'total':(iu+eu)*q}
    def raw_sum(self):
        h=e=0
        for o in self.visible_outputs():
            r=self.raw_row(o); h+=r['hec']; e+=r['ext']
        return h,e,h+e
    def pct(self):
        if self.ratio_mode=='original':
            h,e,t=self.raw_sum(); ip=h/t*100 if t else 0; return ip,100-ip
        return self.internal_pct,self.external_pct
    def row(self,o):
        r=self.raw_row(o)
        # 외주최소화/사용자 입력 비율은 최우선: Total M/H를 내부/외부 비율로 배분
        if self.ratio_mode=='custom':
            ip,ep=self.pct(); r['hec']=r['total']*ip/100; r['ext']=r['total']*ep/100
            return r
        # 외주-종합 Case: 원본 Excel OP1의 Case3 로직을 반영
        # Case3 Internal Unit = 내부 Unit * (1 - GEC/Internal ratio - 종합전환 ratio)
        # Case3 External Unit = 외부 Unit * (1 - 외주감축 ratio) + 내부 Unit * 종합전환 ratio / 배부모수
        # GEC Unit = 내부 Unit * GEC/Internal ratio / 배부모수
        # 프로그램 화면은 GEC 별도 열이 없으므로 GEC M/H를 외부/종합 수행분에 포함하여 Total을 맞춤
        if self.case=='comp':
            q=r['qty']; iu=r['iu']; eu=r['eu']
            if r['part']=='ci':
                gec_ratio=0.30; comp_ratio=0.10; ext_reduce=0.0; denom=0.70
            else:
                gec_ratio=0.00; comp_ratio=0.10; ext_reduce=0.0; denom=0.70
            internal_unit=iu*(1-gec_ratio-comp_ratio)
            gec_unit=iu*gec_ratio/denom if denom else 0
            comp_external_unit=eu*(1-ext_reduce)+iu*comp_ratio/denom if denom else eu
            r['iu']=internal_unit
            r['eu']=gec_unit+comp_external_unit
            r['hec']=round(internal_unit*q)
            r['ext']=round((gec_unit+comp_external_unit)*q)
            r['total']=r['hec']+r['ext']
            r['diff']=str(r.get('diff',''))+' / 외주-종합'
        return r
    def summary(self):
        h=e=t=0
        for o in self.visible_outputs():
            r=self.row(o); h+=r['hec']; e+=r['ext']; t+=r['total']
        return h,e,t
    def recalc(self):
        self.compute_diffs(); self.update_qty()
        if self.ratio_mode=='original':
            h,e,t=self.raw_sum(); self.internal_pct=h/t*100 if t else 0; self.external_pct=100-self.internal_pct
class ScrollFrame(ttk.Frame):
    def __init__(self,parent):
        super().__init__(parent); self.canvas=tk.Canvas(self,bg=CARD,highlightthickness=0); self.inner=ttk.Frame(self.canvas,style='Card.TFrame'); vs=ttk.Scrollbar(self,orient='vertical',command=self.canvas.yview); hs=ttk.Scrollbar(self,orient='horizontal',command=self.canvas.xview); self.canvas.configure(yscrollcommand=vs.set,xscrollcommand=hs.set); self.canvas.create_window((0,0),window=self.inner,anchor='nw'); self.inner.bind('<Configure>',lambda e:self.canvas.configure(scrollregion=self.canvas.bbox('all'))); self.canvas.grid(row=0,column=0,sticky='nsew'); vs.grid(row=0,column=1,sticky='ns'); hs.grid(row=1,column=0,sticky='ew'); self.rowconfigure(0,weight=1); self.columnconfigure(0,weight=1); self.canvas.bind('<Enter>',lambda e:self.canvas.bind_all('<MouseWheel>',self._wheel)); self.canvas.bind('<Leave>',lambda e:self.canvas.unbind_all('<MouseWheel>'))
    def _wheel(self,e): self.canvas.yview_scroll(int(-1*(e.delta/120)),'units')

class StdExcelGrid(ttk.Frame):
    """Excel-like editable standards grid for 산출기준 tabs."""
    def __init__(self, parent, app, part):
        super().__init__(parent)
        self.app=app; self.part=part
        self.canvas=tk.Canvas(self,bg='white',highlightthickness=0)
        self.vs=ttk.Scrollbar(self,orient='vertical',command=self.canvas.yview)
        self.hs=ttk.Scrollbar(self,orient='horizontal',command=self.canvas.xview)
        self.canvas.configure(yscrollcommand=self.vs.set,xscrollcommand=self.hs.set)
        self.canvas.grid(row=0,column=0,sticky='nsew')
        self.vs.grid(row=0,column=1,sticky='ns')
        self.hs.grid(row=1,column=0,sticky='ew')
        self.rowconfigure(0,weight=1); self.columnconfigure(0,weight=1)
        self.edit_cells={}
        self.canvas.bind('<Double-1>', self.on_dbl)
        self.canvas.bind('<MouseWheel>', lambda e:self.canvas.yview_scroll(int(-1*(e.delta/120)),'units'))
        self.canvas.bind('<Shift-MouseWheel>', lambda e:self.canvas.xview_scroll(int(-1*(e.delta/120)),'units'))

    def draw_text(self,x1,y1,x2,y2,text,fill='black',font=('Segoe UI',8),anchor='center',bold=False):
        if bold:
            font=(font[0],font[1],'bold')
        # left aligned text gets slight padding
        if anchor=='w':
            return self.canvas.create_text(x1+6,(y1+y2)//2,text=str(text),fill=fill,font=font,anchor='w',width=max(20,x2-x1-10))
        if anchor=='e':
            return self.canvas.create_text(x2-6,(y1+y2)//2,text=str(text),fill=fill,font=font,anchor='e',width=max(20,x2-x1-10))
        return self.canvas.create_text((x1+x2)//2,(y1+y2)//2,text=str(text),fill=fill,font=font,anchor='center',width=max(20,x2-x1-8))

    def cell(self,x1,y1,x2,y2,text='',bg='white',fg='black',outline='#7F7F7F',font=('Segoe UI',8),anchor='center',bold=False,tag=None):
        self.canvas.create_rectangle(x1,y1,x2,y2,fill=bg,outline=outline,width=1)
        if text!='':
            self.draw_text(x1,y1,x2,y2,text,fg,font,anchor,bold)
        if tag:
            self.edit_cells[tag]=(x1,y1,x2,y2)

    def _wrap_line_count(self, text, width_px, font_size=8):
        # Canvas text wrapping is pixel-based. This approximates required line count so row height
        # expands before drawing and long Guide text is not clipped. Korean/English mixed text is
        # handled conservatively by using a slightly larger character width.
        text=str(text or '')
        if not text:
            return 1
        avg_char_px=max(5.2, font_size*0.72)
        chars=max(8, int((width_px-12)/avg_char_px))
        total=0
        for raw in text.replace('\r','').split('\n'):
            raw=raw.strip()
            if not raw:
                total+=1
            else:
                total+=max(1, (len(raw)+chars-1)//chars)
        return max(1,total)

    def _row_height(self, activity, guide, activity_w, guide_w):
        a_lines=self._wrap_line_count(activity, activity_w, 8)
        g_lines=self._wrap_line_count(guide, guide_w, 8)
        lines=max(a_lines,g_lines,1)
        return max(32, min(150, lines*15+14))

    def render(self):
        self.canvas.delete('all'); self.edit_cells={}
        part=self.part
        if part=='ci':
            entries=list(self.app.m.std['ci'].items())
            cols=[('NO.',45),('Activity',360),('단위',70),('SPI',70),('상',70),('중',70),('하',70),('SPI',70),('상',70),('중',70),('하',70),('Guide',560)]
            title='C&I설계 화공 FEED사업 M/H 산출기준 (Rev.3)'
            version='일반 Ver.'
            int_cols=[3,4,5,6]; ext_cols=[7,8,9,10]
        else:
            entries=list(self.app.m.std['tel'].items())
            cols=[('NO.',45),('Activity',360),('단위',70),('상',80),('중',80),('하',80),('상',80),('중',80),('하',80),('Guide',560)]
            title='통신설계 화공 FEED사업 M/H 산출기준 (Rev.3)'
            version='일반 Ver.'
            int_cols=[3,4,5]; ext_cols=[6,7,8]
        xs=[0]
        for _,w in cols: xs.append(xs[-1]+w)
        W=xs[-1]
        y=0
        # top title and version/notes, close to Excel original
        self.cell(0,y,W,y+34,title,bg='#FFFFFF',fg='black',outline='#FFFFFF',font=('Malgun Gothic',13),bold=True)
        self.draw_text(W-140,y,W-8,y+34,'2026-08-11',fill='black',font=('Malgun Gothic',8),anchor='e',bold=True)
        y+=34
        self.cell(0,y,360,y+44,'산출기준 적용 Version\n(일반 vs 외주최소화 Ver.)',bg='#BFE3F0',fg='black',outline='black',font=('Malgun Gothic',9),bold=True)
        self.cell(360,y,720,y+44,version,bg='#FFC000',fg='red',outline='black',font=('Malgun Gothic',12),bold=True)
        self.cell(720,y,W,y+44,'Notes\n1. Project 표준 난이도, Unit M/H 난이도 및 SPI 적용 여부를 기준으로 M/H를 산정합니다.\n2. 노란색/주황색 숫자 셀은 더블클릭으로 직접 조정 가능하며 Summary/Output에 즉시 반영됩니다.',bg='#F2F2F2',fg='black',outline='#FFFFFF',font=('Malgun Gothic',8),anchor='w'); y+=54
        # group header section
        self.cell(xs[0],y,xs[3],y+72,'',bg='#DDEBF7',outline='black')
        self.cell(xs[0],y,xs[1],y+72,'NO.',bg='#DDEBF7',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[1],y,xs[2],y+72,'Activity',bg='#DDEBF7',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[2],y,xs[3],y+72,'단위',bg='#DDEBF7',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[int_cols[0]],y,xs[int_cols[-1]+1],y+24,'난이도별 Unit M/H',bg='#FCE4D6',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[int_cols[0]],y+24,xs[int_cols[-1]+1],y+48,'내부 Unit M/H',bg='#FCE4D6',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[ext_cols[0]],y,xs[ext_cols[-1]+1],y+24,'난이도별 Unit M/H',bg='#FCE4D6',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[ext_cols[0]],y+24,xs[ext_cols[-1]+1],y+48,'외주 Unit M/H',bg='#FCE4D6',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[ext_cols[-1]+1],y,xs[-1],y+48,'MAN-HOUR 산출 지침',bg='#BFE3F0',outline='black',font=('Malgun Gothic',9),bold=True)
        # sub headers
        for i,(name,w) in enumerate(cols):
            if i>=3 and i<ext_cols[-1]+1:
                self.cell(xs[i],y+48,xs[i+1],y+72,name,bg='#FCE4D6',outline='black',font=('Malgun Gothic',8),bold=True)
        self.cell(xs[-2],y+48,xs[-1],y+72,'Guide',bg='#BFE3F0',outline='black',font=('Malgun Gothic',8),bold=True)
        y+=72
        # rows: each row height is calculated from Activity/Guide text length to prevent clipping.
        for idx,(key,r) in enumerate(entries,1):
            activity=r.get('activity','')
            guide=r.get('guide','')
            row_h=self._row_height(activity, guide, xs[2]-xs[1], xs[-1]-xs[-2])
            bg='#FFFFFF' if idx%2 else '#F7FBFF'
            self.cell(xs[0],y,xs[1],y+row_h,idx,bg=bg,outline='#B7B7B7',font=('Malgun Gothic',8))
            self.cell(xs[1],y,xs[2],y+row_h,activity,bg=bg,outline='#B7B7B7',font=('Malgun Gothic',8),anchor='w')
            self.cell(xs[2],y,xs[3],y+row_h,r.get('unit',''),bg=bg,outline='#B7B7B7',font=('Malgun Gothic',8))
            # values
            if part=='ci':
                values=[('int','SPI',r.get('int',{}).get('SPI',0)),('int','상',r.get('int',{}).get('상',0)),('int','중',r.get('int',{}).get('중',0)),('int','하',r.get('int',{}).get('하',0)),('ext','SPI',r.get('ext',{}).get('SPI',0)),('ext','상',r.get('ext',{}).get('상',0)),('ext','중',r.get('ext',{}).get('중',0)),('ext','하',r.get('ext',{}).get('하',0))]
                start=3
            else:
                values=[('int','상',r.get('int',{}).get('상',0)),('int','중',r.get('int',{}).get('중',0)),('int','하',r.get('int',{}).get('하',0)),('ext','상',r.get('ext',{}).get('상',0)),('ext','중',r.get('ext',{}).get('중',0)),('ext','하',r.get('ext',{}).get('하',0))]
                start=3
            for j,(typ,diff,val) in enumerate(values):
                c=start+j
                val_bg='#FFF2CC' if typ=='int' else '#E2F0D9'
                tag=(part,key,typ,diff)
                self.cell(xs[c],y,xs[c+1],y+row_h,val,bg=val_bg,outline='#B7B7B7',font=('Malgun Gothic',8),anchor='e',tag=tag)
            self.cell(xs[-2],y,xs[-1],y+row_h,guide,bg=bg,outline='#B7B7B7',font=('Malgun Gothic',8),anchor='w')
            y+=row_h
        self.canvas.configure(scrollregion=(0,0,W,y+20))

    def on_dbl(self,event):
        x=self.canvas.canvasx(event.x); y=self.canvas.canvasy(event.y)
        target=None
        for tag,(x1,y1,x2,y2) in self.edit_cells.items():
            if x1<=x<=x2 and y1<=y<=y2:
                target=tag; box=(x1,y1,x2,y2); break
        if not target: return
        part,key,typ,diff=target; x1,y1,x2,y2=box
        old=str(self.app.m.std[part][key].get(typ,{}).get(diff,0))
        entry=ttk.Entry(self.canvas)
        win=self.canvas.create_window(x1,y1,window=entry,anchor='nw',width=x2-x1,height=y2-y1)
        entry.insert(0,old); entry.focus_set(); entry.select_range(0,'end')
        def commit(event=None):
            val=entry.get().strip()
            try:
                n=float(val.replace(',','')) if val!='' else 0.0
            except Exception:
                messagebox.showerror('입력 오류','숫자만 입력할 수 있습니다.')
                return
            self.canvas.delete(win)
            self.app.update_std_value(part,key,typ,diff,n)
        def cancel(event=None):
            self.canvas.delete(win)
        entry.bind('<Return>',commit); entry.bind('<FocusOut>',commit); entry.bind('<Escape>',cancel)


class OpExcelGrid(ttk.Frame):
    """Excel-like OP sheet grid with grouped headers and tree-style section rows."""
    def __init__(self, parent, app, sheet):
        super().__init__(parent)
        self.app=app; self.sheet=sheet
        self.canvas=tk.Canvas(self,bg='white',highlightthickness=0)
        self.vs=ttk.Scrollbar(self,orient='vertical',command=self.canvas.yview)
        self.hs=ttk.Scrollbar(self,orient='horizontal',command=self.canvas.xview)
        self.canvas.configure(yscrollcommand=self.vs.set,xscrollcommand=self.hs.set)
        self.canvas.grid(row=0,column=0,sticky='nsew'); self.vs.grid(row=0,column=1,sticky='ns'); self.hs.grid(row=1,column=0,sticky='ew')
        self.rowconfigure(0,weight=1); self.columnconfigure(0,weight=1)
        self.canvas.bind('<MouseWheel>', lambda e:self.canvas.yview_scroll(int(-1*(e.delta/120)),'units'))
        self.canvas.bind('<Shift-MouseWheel>', lambda e:self.canvas.xview_scroll(int(-1*(e.delta/120)),'units'))

    def draw_text(self,x1,y1,x2,y2,text,fill='black',font=('Malgun Gothic',8),anchor='center',bold=False):
        if bold: font=(font[0],font[1],'bold')
        width=max(20,x2-x1-8)
        if anchor=='w': return self.canvas.create_text(x1+6,(y1+y2)//2,text=str(text),fill=fill,font=font,anchor='w',width=width)
        if anchor=='e': return self.canvas.create_text(x2-6,(y1+y2)//2,text=str(text),fill=fill,font=font,anchor='e',width=width)
        return self.canvas.create_text((x1+x2)//2,(y1+y2)//2,text=str(text),fill=fill,font=font,anchor='center',width=width)

    def cell(self,x1,y1,x2,y2,text='',bg='white',fg='black',outline='#7F7F7F',font=('Malgun Gothic',8),anchor='center',bold=False,width=1):
        self.canvas.create_rectangle(x1,y1,x2,y2,fill=bg,outline=outline,width=width)
        if text!='': self.draw_text(x1,y1,x2,y2,text,fg,font,anchor,bold)

    def money(self,v,d=0):
        try: return f"{float(v):,.{d}f}"
        except Exception: return str(v)

    def section_name(self, code):
        try:
            n=int(str(code).replace('A',''))
        except Exception:
            return 'OTHER'
        if n<=6: return '1  GENERAL'
        if n<=11: return '2  SPECIFICATION'
        if n<=14: return '3  CALCULATION'
        if n<=20: return '4  구매관련 / DATA SHEET / MR & TBE'
        if n<=31: return '5  DRAWING / INDEX / INTERFACE'
        if n<=33: return '6  MTO & 타부서 INFORM'
        return '7  OTHERS'

    def rows_for_sheet(self):
        m=self.app.m
        rows=[]
        old_case=m.case
        if self.sheet=='OP2-단종': m.case='short'
        elif self.sheet=='OP2-종합': m.case='comp'
        # OP sheets use current part selection. If 전체, include both.
        for o in m.visible_outputs():
            rows.append(m.row(o))
        m.case=old_case
        return rows

    def render(self):
        self.canvas.delete('all')
        m=self.app.m
        rows=self.rows_for_sheet()
        title='C&I / Telecom FEED사업 M/H 산출서 (Rev.3)'
        if m.part=='ci': title='C&I설계 화공 FEED사업 M/H 산출서 (Rev.3)'
        elif m.part=='tel': title='통신설계 화공 FEED사업 M/H 산출서 (Rev.3)'
        ver='[ 외주최소화 Ver. ]' if m.external_min=='Yes' else '[ 일반 Ver. ]'
        # Excel-like columns
        if self.sheet=='OP1':
            cols=[('No.',45),('Part',70),('Activity',340),('단위',60),('수량',70),('난이도',65),('내부 Unit',80),('외주 Unit',80),('내부 M/H',90),('외주 M/H',90),('Total M/H',95),('Remarks',220)]
        else:
            cols=[('No.',45),('Part',70),('Activity',360),('단위',60),('수량',70),('난이도',75),('내부 M/H',95),('외주 M/H',95),('Total M/H',100),('Remarks',260)]
        xs=[0]
        for _,w in cols: xs.append(xs[-1]+w)
        W=xs[-1]
        y=0
        # Title / notes
        self.cell(0,y,W,y+32,title,bg='white',outline='white',font=('Malgun Gothic',13),bold=True)
        self.draw_text(W-145,y,W-8,y+32,'2026-08-11',font=('Malgun Gothic',8),anchor='e',bold=True); y+=32
        self.cell(0,y,W,y+30,ver,bg='white',fg='red',outline='white',font=('Malgun Gothic',12),bold=True); y+=36
        # base data summary and notes
        self.cell(0,y,470,y+56,'산출 Base Data Summary',bg='white',outline='white',font=('Malgun Gothic',8),anchor='w',bold=True)
        if m.part in ('both','ci'):
            self.cell(0,y+20,260,y+38,'Total ICSS (DCS+ESD+F&G)+MMS IO 수량',bg='white',outline='black',anchor='w',bold=True)
            self.cell(260,y+20,340,y+38,self.money(m.nval('ci','A01_QTY')),bg='#40D050',outline='black',anchor='e',bold=True)
            self.cell(340,y+20,470,y+38,'Actual수량 기입',bg='white',fg='red',outline='black',anchor='w')
            self.cell(0,y+38,260,y+56,'Total Instrument 수량',bg='white',outline='black',anchor='w',bold=True)
            self.cell(260,y+38,340,y+56,self.money(m.nval('ci','A02_QTY')),bg='#40D050',outline='black',anchor='e',bold=True)
            self.cell(340,y+38,470,y+56,'F&G Inst. 포함',bg='white',outline='black',anchor='w')
        elif m.part=='tel':
            self.cell(0,y+20,260,y+38,'Telecom System 수량',bg='white',outline='black',anchor='w',bold=True)
            self.cell(260,y+20,340,y+38,self.money(m.nval('tel','A01_QTY')),bg='#40D050',outline='black',anchor='e',bold=True)
            self.cell(340,y+20,470,y+38,'System Qty',bg='white',outline='black',anchor='w')
        self.cell(770,y,1180,y+56,'Notes\n1. FEED 산출 기준은 산출기준 탭의 Unit M/H와 Project Condition을 기준으로 적용합니다.\n2. 섹션별 트리 구조로 Activity를 구분하고, OP1/OP2 결과를 동일 화면에서 검토합니다.',bg='#D9D9D9',outline='white',font=('Malgun Gothic',8),anchor='w')
        y+=66
        self.cell(0,y,300,y+22,'PROJECT:',bg='#63D663',outline='black',font=('Malgun Gothic',8),anchor='w',bold=True)
        self.cell(300,y,W,y+22,str(m.project or ''),bg='#9DE28E',outline='black',font=('Malgun Gothic',8),anchor='w')
        y+=22
        # grouped headers
        self.cell(0,y,W,y+20,'MAN-HOUR 산출 TABLE',bg='#8BD56B',outline='black',font=('Malgun Gothic',9),bold=True); y+=20
        self.cell(xs[0],y,xs[6],y+60,'기본 정보',bg='#DDEBF7',outline='black',bold=True)
        if self.sheet=='OP1':
            self.cell(xs[6],y,xs[8],y+28,'Unit M/H',bg='#BFE3F0',outline='black',bold=True)
            self.cell(xs[8],y,xs[10],y+28,'M/H',bg='#C6E0B4',outline='black',bold=True)
            self.cell(xs[10],y,xs[11],y+60,'Total M/H',bg='#BFE3F0',outline='black',bold=True)
            self.cell(xs[11],y,xs[12],y+60,'Remarks',bg='#EAD1EA',outline='black',bold=True)
        else:
            self.cell(xs[6],y,xs[9],y+28,'M/H',bg='#C6E0B4',outline='black',bold=True)
            self.cell(xs[9],y,xs[10],y+60,'Remarks',bg='#EAD1EA',outline='black',bold=True)
        y+=28
        for i,(name,w) in enumerate(cols):
            if self.sheet=='OP1' or i<10:
                self.cell(xs[i],y,xs[i+1],y+32,name,bg='#DDEBF7' if i<6 else ('#BFE3F0' if 'Unit' in name or 'Total' in name else '#C6E0B4'),outline='black',font=('Malgun Gothic',8),bold=True)
        y+=32
        # body, section rows
        current=None; no=1
        totals={'hec':0,'ext':0,'total':0}
        row_h=28
        for r in rows:
            sec=self.section_name(r.get('code',''))
            if sec!=current:
                current=sec
                self.cell(0,y,W,y+28,current,bg='#DDEBF7',outline='black',font=('Malgun Gothic',9),anchor='w',bold=True)
                y+=28
            bg='#FFFFFF' if no%2 else '#F7FBFF'
            vals=[no, self.app.part_label(r.get('part')), r.get('activity',''), r.get('unit',''), self.money(r.get('qty'),1), r.get('diff','')]
            if self.sheet=='OP1':
                vals += [self.money(r.get('iu'),2), self.money(r.get('eu'),2), self.money(r.get('hec')), self.money(r.get('ext')), self.money(r.get('total')), '']
            else:
                vals += [self.money(r.get('hec')), self.money(r.get('ext')), self.money(r.get('total')), '']
            for i,v in enumerate(vals):
                anchor='w' if i in (2,11) else 'e' if i>=4 and i not in (5,) else 'center'
                cell_bg=bg
                if self.sheet=='OP1' and i in (6,7): cell_bg='#FFF2CC'
                if (self.sheet=='OP1' and i in (8,9)) or (self.sheet!='OP1' and i in (6,7)): cell_bg='#E2F0D9'
                if (self.sheet=='OP1' and i==10) or (self.sheet!='OP1' and i==8): cell_bg='#DDEBF7'
                self.cell(xs[i],y,xs[i+1],y+row_h,v,bg=cell_bg,outline='#B7B7B7',font=('Malgun Gothic',8),anchor=anchor,bold=(('Total' in cols[i][0]) or (self.sheet=='OP1' and i==10) or (self.sheet!='OP1' and i==8)))
            totals['hec']+=float(r.get('hec') or 0); totals['ext']+=float(r.get('ext') or 0); totals['total']+=float(r.get('total') or 0)
            no+=1; y+=row_h
        # summary rows
        self.cell(0,y,W,y+28,'SUMMARY',bg='#FFFF00',outline='black',font=('Malgun Gothic',8),anchor='w',bold=True)
        if self.sheet=='OP1':
            self.cell(xs[8],y,xs[9],y+28,self.money(totals['hec']),bg='#FFFF00',outline='black',anchor='e',bold=True)
            self.cell(xs[9],y,xs[10],y+28,self.money(totals['ext']),bg='#FFFF00',outline='black',anchor='e',bold=True)
            self.cell(xs[10],y,xs[11],y+28,self.money(totals['total']),bg='#FFFF00',outline='black',anchor='e',bold=True)
        else:
            self.cell(xs[6],y,xs[7],y+28,self.money(totals['hec']),bg='#FFFF00',outline='black',anchor='e',bold=True)
            self.cell(xs[7],y,xs[8],y+28,self.money(totals['ext']),bg='#FFFF00',outline='black',anchor='e',bold=True)
            self.cell(xs[8],y,xs[9],y+28,self.money(totals['total']),bg='#FFFF00',outline='black',anchor='e',bold=True)
        y+=28
        mm=totals['total']/m.base_mh if m.base_mh else 0
        avg=mm/m.months if m.months else 0
        self.cell(0,y,W,y+24,f"Base M/H: {self.money(m.base_mh)}   |   Total M/M: {self.money(mm,1)}   |   설계기간: {self.money(m.months)} 개월   |   평균 투입인원: {self.money(avg,1)} 명/month",bg='white',outline='black',font=('Malgun Gothic',8),anchor='w',bold=True)
        y+=30
        self.canvas.configure(scrollregion=(0,0,W,y+20))

class App(tk.Tk):
    def __init__(self):
        super().__init__(); self.m=Model(); self.title('C&I / Telecom FEED M/H Calculator V.0.4 / 2026-08-11'); self.geometry('1680x950'); self.configure(bg=BG); self.vars={}; self.lock=False; self.std_item_map={}
        try:
            icon_path=Path(__file__).with_name('feed_mh_calculator_icon.ico')
            if icon_path.exists(): self.iconbitmap(str(icon_path))
        except Exception:
            pass
        self.style(); self.build(); self.load_user_state(); self.refresh_all(); self.protocol('WM_DELETE_WINDOW', self.on_close)
    def style(self):
        s=ttk.Style(); s.theme_use('clam')
        s.configure('TFrame',background=BG); s.configure('Card.TFrame',background=CARD)
        s.configure('TLabel',background=BG,foreground=TEXT,font=('Segoe UI',9)); s.configure('Card.TLabel',background=CARD,foreground=TEXT,font=('Segoe UI',9))
        s.configure('Title.TLabel',background=PRIMARY,foreground='white',font=('Segoe UI',17,'bold'),padding=15)
        s.configure('TButton',font=('Segoe UI',9,'bold'),padding=8); s.configure('Accent.TButton',background=ACCENT,foreground='white')
        s.map('Accent.TButton',background=[('active','#0B5CAB')],foreground=[('active','white')])
        s.configure('TEntry',padding=5); s.configure('TCombobox',padding=5)
        s.configure('TNotebook',background=BG,tabmargins=[2,4,2,0]); s.configure('TNotebook.Tab',font=('Segoe UI',9,'bold'),padding=[10,5],background='#DFE7EF')
        s.map('TNotebook.Tab',background=[('selected','#FFFFFF')],foreground=[('selected',PRIMARY)])
        s.configure('Treeview',font=('Segoe UI',9),rowheight=30,background='white',fieldbackground='white',bordercolor=LINE,relief='flat')
        s.configure('Treeview.Heading',font=('Segoe UI',9,'bold'),background=HEAD,foreground='white',padding=6)
    def build(self):
        ttk.Label(self,text='C&I / Telecom FEED M/H Calculator V.0.4 / 2026-08-11',style='Title.TLabel').pack(fill='x')
        body=ttk.Frame(self); body.pack(fill='both',expand=True,padx=12,pady=12)
        # Left Master Control panel is scrollable so lower controls are never clipped on small screens.
        left_holder=ttk.Frame(body,style='Card.TFrame')
        left_holder.pack(side='left',fill='y',padx=(0,12))
        left_holder.pack_propagate(False)
        left_holder.configure(width=410)
        left_scroll=ScrollFrame(left_holder)
        left_scroll.pack(fill='both',expand=True)
        left=ttk.Frame(left_scroll.inner,style='Card.TFrame',padding=14)
        left.pack(fill='both',expand=True)
        right=ttk.Frame(body); right.pack(side='right',fill='both',expand=True)
        self._make_logo_widget(left)
        ttk.Label(left,text='MASTER CONTROL',style='Card.TLabel',font=('Segoe UI',14,'bold')).pack(anchor='w'); ttk.Label(left,text='왼쪽 Master 값을 기준으로 Input 공통항목과 Output이 자동 연동됩니다.',style='Card.TLabel',wraplength=280,foreground='#506980').pack(anchor='w',pady=(2,8))
        def combo(label,key,vals,cmd=None):
            ttk.Label(left,text=label,style='Card.TLabel').pack(anchor='w',pady=(8,2)); self.vars[key]=tk.StringVar(value=vals[0]); cb=ttk.Combobox(left,textvariable=self.vars[key],values=vals,state='readonly'); cb.pack(fill='x');
            if cmd: cb.bind('<<ComboboxSelected>>',cmd)
        def entry(label,key,val,cmd=None):
            ttk.Label(left,text=label,style='Card.TLabel').pack(anchor='w',pady=(8,2)); self.vars[key]=tk.StringVar(value=str(val)); e=ttk.Entry(left,textvariable=self.vars[key]); e.pack(fill='x');
            if cmd: e.bind('<KeyRelease>',cmd); e.bind('<FocusOut>',cmd); e.bind('<Return>',cmd)
        combo('Part','part',['전체','CI (계장)','TEL (통신)'],lambda e:self.apply_left()); combo('Case','case',['외주-단종','외주-종합'],lambda e:self.apply_left()); combo('비율 적용 방식','ratio',['원본 기준','사용자 입력'],lambda e:self.apply_left()); combo('외주최소화','external',['No','Yes'],self.on_external)
        entry('Project','project',''); entry('Base M/H per 1 M/M','base',161); entry('설계기간 개월','months',4); entry('외주-단종 단가','rate_short',32000); entry('외주-종합 단가','rate_comp',39000); entry('내부 %','ip','',self.on_ip_change); entry('외부 %','ep','',self.on_ep_change)
        ttk.Button(left,text='Output 재계산',style='Accent.TButton',command=self.apply_left).pack(fill='x',pady=(12,4))
        ttk.Button(left,text='입력값 저장',style='Accent.TButton',command=self.save_user_state).pack(fill='x',pady=(4,4))
        ttk.Button(left,text='저장값 불러오기',style='Accent.TButton',command=lambda:self.load_user_state(show_msg=True)).pack(fill='x',pady=(4,4))
        ttk.Button(left,text='입력값 내보내기',command=self.export_user_state).pack(fill='x',pady=(4,4))
        ttk.Button(left,text='입력값 가져오기',command=self.import_user_state).pack(fill='x',pady=(4,4))
        ttk.Button(left,text='초기화',command=self.reset_user_state).pack(fill='x',pady=(4,4))
        ttk.Button(left,text='Word Report 생성',style='Accent.TButton',command=self.generate_report).pack(fill='x',pady=(12,4))
        ttk.Label(left,text='외주최소화는 왼쪽 Master Control에서만 변경합니다. Input 수정 탭의 파란색 Master연동 행은 왼쪽 값이 그대로 표시됩니다. 초록색 자동결과는 선택값에 따라 즉시 계산되어 표시됩니다.',style='Card.TLabel',wraplength=280,foreground='#7A4F00').pack(fill='x',pady=12)
        self.tabs=ttk.Notebook(right); self.tabs.pack(fill='both',expand=True)
        self.input_tab=ttk.Frame(self.tabs); self.tabs.add(self.input_tab,text='Input 수정'); self.input_frame=ScrollFrame(self.input_tab); self.input_frame.pack(fill='both',expand=True)
        self.guide_tab=ttk.Frame(self.tabs); self.tabs.add(self.guide_tab,text='Guide / Help')
        self.guide_frame=ScrollFrame(self.guide_tab); self.guide_frame.pack(fill='both',expand=True)
        self.render_guide_tab()
        self.trees={}
        self.summary_tab=ttk.Frame(self.tabs); self.tabs.add(self.summary_tab,text='Summary')
        self.summary_area=ttk.Frame(self.summary_tab,style='Card.TFrame',padding=14); self.summary_area.pack(fill='both',expand=True)
        self.output_ci_tab=ttk.Frame(self.tabs); self.tabs.add(self.output_ci_tab,text='Output_CI')
        self.output_ci_frame=ScrollFrame(self.output_ci_tab); self.output_ci_frame.pack(fill='both',expand=True)
        self.output_tel_tab=ttk.Frame(self.tabs); self.tabs.add(self.output_tel_tab,text='Output_TEL')
        self.output_tel_frame=ScrollFrame(self.output_tel_tab); self.output_tel_frame.pack(fill='both',expand=True)
        self.op1_tab=ttk.Frame(self.tabs); self.tabs.add(self.op1_tab,text='OP1')
        self.op_grid_op1=OpExcelGrid(self.op1_tab,self,'OP1'); self.op_grid_op1.pack(fill='both',expand=True)
        self.op2_short_tab=ttk.Frame(self.tabs); self.tabs.add(self.op2_short_tab,text='OP2-단종')
        self.op_grid_short=OpExcelGrid(self.op2_short_tab,self,'OP2-단종'); self.op_grid_short.pack(fill='both',expand=True)
        self.op2_comp_tab=ttk.Frame(self.tabs); self.tabs.add(self.op2_comp_tab,text='OP2-종합')
        self.op_grid_comp=OpExcelGrid(self.op2_comp_tab,self,'OP2-종합'); self.op_grid_comp.pack(fill='both',expand=True)
        self.std_tab_ci=ttk.Frame(self.tabs); self.tabs.add(self.std_tab_ci,text='산출기준_CI')
        self.std_grid_ci=StdExcelGrid(self.std_tab_ci,self,'ci'); self.std_grid_ci.pack(fill='both',expand=True)
        self.std_tab_tel=ttk.Frame(self.tabs); self.tabs.add(self.std_tab_tel,text='산출기준_TEL')
        self.std_grid_tel=StdExcelGrid(self.std_tab_tel,self,'tel'); self.std_grid_tel.pack(fill='both',expand=True)
    def make_tree(self,parent,cols):
        fr=ttk.Frame(parent); fr.pack(fill='both',expand=True); t=ttk.Treeview(fr,columns=cols,show='headings'); vs=ttk.Scrollbar(fr,orient='vertical',command=t.yview); hs=ttk.Scrollbar(fr,orient='horizontal',command=t.xview); t.configure(yscrollcommand=vs.set,xscrollcommand=hs.set); t.grid(row=0,column=0,sticky='nsew'); vs.grid(row=0,column=1,sticky='ns'); hs.grid(row=1,column=0,sticky='ew'); fr.rowconfigure(0,weight=1); fr.columnconfigure(0,weight=1)
        for c in cols:
            t.heading(c,text=c)
            # Excel처럼 열 폭을 고정하고, 창 크기에 따라 숫자열이 찌그러지지 않도록 stretch를 끕니다.
            t.column(c,width=120,stretch=False,anchor='center' if c not in ['Activity','Guide','Remarks'] else 'w')
        t.tag_configure('odd', background='#F7FAFD'); t.tag_configure('even', background='white')
        t.tag_configure('std_ci_even', background='#EAF5FF'); t.tag_configure('std_ci_odd', background='#FFFFFF')
        t.tag_configure('std_tel_even', background='#EAF8F0'); t.tag_configure('std_tel_odd', background='#FFFFFF')
        t.tag_configure('locked', background='#F2F2F2')
        t.bind('<MouseWheel>',lambda e,tr=t: tr.yview_scroll(int(-1*(e.delta/120)),'units'))
        t.bind('<Shift-MouseWheel>',lambda e,tr=t: tr.xview_scroll(int(-1*(e.delta/120)),'units'))
        t.bind('<Double-1>', self.on_tree_double_click)
        return t
    def on_external(self,e=None):
        if self.vars['external'].get()=='Yes':
            self.vars['ratio'].set('사용자 입력'); self.vars['ip'].set('75.0'); self.vars['ep'].set('25.0')
        self.apply_left()
    def on_ip_change(self,e=None):
        if self.lock: return
        v=self.vars['ip'].get();
        try:
            val=float(v); self.lock=True; self.vars['ep'].set(f'{100-val:.1f}'); self.vars['ratio'].set('사용자 입력'); self.lock=False
        except: pass
    def on_ep_change(self,e=None):
        if self.lock: return
        v=self.vars['ep'].get();
        try:
            val=float(v); self.lock=True; self.vars['ip'].set(f'{100-val:.1f}'); self.vars['ratio'].set('사용자 입력'); self.lock=False
        except: pass
    def auto_width(self,t,rows):
        # 산출기준 탭은 Excel 원본처럼 고정 폭 레이아웃을 사용합니다.
        # 값 수정 후 재렌더링 시 Treeview가 전체 폭에 맞춰 숫자열을 압축하는 문제를 방지합니다.
        if hasattr(self,'trees') and t in (self.trees.get('산출기준_CI'), self.trees.get('산출기준_TEL')):
            if t == self.trees.get('산출기준_CI'):
                widths=[360,70,78,78,78,78,78,78,78,78,620]
                anchors=['w','center','e','e','e','e','e','e','e','e','w']
            else:
                widths=[360,70,85,85,85,85,85,85,620]
                anchors=['w','center','e','e','e','e','e','e','w']
            for c,w,a in zip(t['columns'],widths,anchors):
                t.column(c,width=w,minwidth=w,stretch=False,anchor=a)
            return
        for idx,c in enumerate(t['columns']):
            ml=len(c)
            for r in rows[:500]:
                if idx<len(r): ml=max(ml,len(str(r[idx])))
            width=max(105,min(620,ml*8+45))
            if any(k in c for k in ['Unit','단위','난이도','Scale','Qty','M/H','Total']):
                width=max(width,115)
            if any(k in c for k in ['Activity','Guide','Remarks']) or ml>32:
                width=max(width,280)
            t.column(c,width=width,minwidth=min(width,160),stretch=False,anchor='w' if c in ['Activity','Guide','Remarks'] else 'center')
    def part_label(self,p): return {'ci':'CI (계장)','tel':'TEL (통신)','both':'전체'}.get(p,p)
    def apply_left(self):
        m=self.m; m.part=REV[self.vars['part'].get()]; m.case='short' if self.vars['case'].get()=='외주-단종' else 'comp'; m.ratio_mode='original' if self.vars['ratio'].get()=='원본 기준' else 'custom'; m.external_min=self.vars['external'].get(); m.project=self.vars['project'].get(); m.base_mh=num(self.vars['base'].get()); m.months=num(self.vars['months'].get()); m.rate_short=num(self.vars['rate_short'].get()); m.rate_comp=num(self.vars['rate_comp'].get())
        if m.external_min=='Yes': m.ratio_mode='custom'
        if m.ratio_mode=='custom': m.internal_pct=num(self.vars['ip'].get()); m.external_pct=100-m.internal_pct; self.vars['ep'].set(f1(m.external_pct))
        m.sync_common(); m.recalc(); self.refresh_all(); self.save_user_state(silent=True)
    def refresh_all(self):
        self.m.sync_common(); self.m.recalc();
        if self.m.ratio_mode=='original': self.lock=True; self.vars['ip'].set(f1(self.m.internal_pct)); self.vars['ep'].set(f1(self.m.external_pct)); self.lock=False
        self.render_input_editor(); self.render_summary(); self.render_output_form('ci'); self.render_output_form('tel'); self.render_op1(); self.render_op2('OP2-단종'); self.render_op2('OP2-종합'); self.render_std()
    def render_guide_tab(self):
        area=self.guide_frame.inner
        for w in area.winfo_children(): w.destroy()
        def title(txt,r):
            tk.Label(area,text=txt,bg=CARD,fg='black',font=('Segoe UI',16,'bold'),anchor='center',padx=8,pady=8).grid(row=r,column=0,columnspan=2,sticky='nsew',padx=1,pady=1)
        def section(txt,r):
            tk.Label(area,text=txt,bg=CARD,fg='black',font=('Segoe UI',12,'bold'),anchor='w',padx=8,pady=8).grid(row=r,column=0,columnspan=2,sticky='nsew',padx=1,pady=(10,1))
        def block(r,head,body='',warn=False):
            bg='#FFFFFF'
            fg='#C00000' if warn else TEXT
            f=tk.Frame(area,bg=bg,highlightbackground=LINE,highlightthickness=1)
            f.grid(row=r,column=0,columnspan=2,sticky='nsew',padx=22,pady=4)
            tk.Label(f,text=head,bg=bg,fg=PRIMARY,font=('Segoe UI',10,'bold'),anchor='w',padx=10,pady=7).pack(fill='x')
            if body:
                for line in body.split('\n'):
                    tk.Label(f,text='- '+line,bg=bg,fg=fg,font=('Segoe UI',9),anchor='w',justify='left',wraplength=1060,padx=24,pady=2).pack(fill='x')
        r=0
        title('M/H 산출 가이드',r); r+=1
        tk.Label(area,text='첨부 #3-1  |  Standalone Program Guide',bg=CARD,fg='#555555',font=('Segoe UI',9),anchor='e',padx=12,pady=2).grid(row=r,column=0,columnspan=2,sticky='nsew'); r+=1
        section('I. 각 시트(탭) 관련 안내',r); r+=1
        items=[
            ('1. Summary 시트','본 프로그램의 Master Control 입력값 기준으로 C&I + 통신 M/H Summary를 표시합니다.\nPart, Case, 외주최소화, 내부/외부 비율, Base M/H, 설계기간, 단가가 Summary와 Output에 공통 반영됩니다.\n외주 예가, Total M/H, Total M/M, 평균 투입인원 확인용입니다.'),
            ('2. 산출기준_CI / 산출기준_TEL 시트','내부/외주 Unit M/H 및 난이도별 산출 기준을 확인하는 기준표입니다.\nActivity, 단위, 난이도, 내부/외부 Unit M/H, Guide 내용을 검토하는 용도입니다.\n열 폭은 실제 셀 내용에 맞춰 자동 확대되어 긴 설명도 확인할 수 있습니다.'),
            ('3. Input 수정 시트','M/H 산출을 위한 주요 입력값과 Project Condition을 입력합니다.\n노란색은 필수 입력, 파란색은 Master Control 연동, 초록색은 자동결과, 보라색은 선택 입력입니다.\n선택형 항목은 기존 v16 방식대로 Combobox에서 선택할 수 있습니다.'),
            ('4. Output_CI / Output_TEL 시트','Part별 산출서 형식으로 Activity, Unit, Qty, 난이도, 내부/외부 M/H, Total M/H를 표시합니다.\n좌측 Master Control의 Case와 외주최소화 조건이 반영됩니다.'),
            ('5. OP1 / OP2 시트','OP1은 Activity별 내부/외부 Unit 및 M/H 상세 확인용입니다.\nOP2-단종 및 OP2-종합은 보고용 M/H 내역 확인용입니다.'),
            ('6. Word Report 생성','현재 화면의 Master Control, Summary, 주요 산출 정보를 기준으로 Word Report를 생성합니다.\n생성 후 저장 폴더가 자동으로 열립니다.')]
        for h,b in items:
            block(r,h,b); r+=1
        section('II. 각 Case 관련 안내',r); r+=1
        for h,b in [
            ('1. 외주-단종 Case','외부 수행분을 외주-단종 단가로 예가 산정합니다.\n내부 검토와 외주 수행을 구분하는 일반 FEED 산출에 사용합니다.'),
            ('2. 외주-종합 Case','외부 수행분을 외주-종합 단가로 예가 산정합니다.\n외주 업체가 복수 업무를 통합 수행하는 조건 검토에 사용합니다.')]:
            block(r,h,b); r+=1
        section('III. 외주최소화 적용 안내',r); r+=1
        for h,b,w in [
            ('1. 외주최소화 Yes','내부/외부 M/H 비율을 75/25 기준으로 적용합니다.\nSummary, Output, OP, Word Report에 동일하게 반영됩니다.',False),
            ('2. 외주최소화 No','원본 산출기준의 내부/외주 Unit M/H 산정 결과를 기준으로 비율을 계산합니다.',False),
            ('3. 원본 기준과 외주최소화','원본 기준은 원본 Unit M/H 구조를 유지하는 방식입니다.\n외주최소화를 적용하려면 사용자 입력 비율을 사용합니다.',True)]:
            block(r,h,b,w); r+=1
        section('IV. 사용 순서',r); r+=1
        block(r,'1. 기본 정보 입력','좌측 Master Control에서 Project, Part, Case, Base M/H, 설계기간, 외주 단가를 입력합니다.'); r+=1
        block(r,'2. Input 수정','Input 수정 탭에서 수량 및 선택 조건을 검토하고 필요한 값을 수정합니다.'); r+=1
        block(r,'3. Output 확인 및 Word Report 생성','Summary, Output, OP 탭에서 결과를 확인한 후 Word Report를 생성합니다.'); r+=1
        area.columnconfigure(0,minsize=260)
        area.columnconfigure(1,minsize=900,weight=1)

    def fallback_diff(self, part, code):
        # Display-side fallback so computed difficulty cells are never blank.
        val=self.m.val(part,code)
        if val not in (None,''):
            return val
        if not code.endswith('_DIFF'):
            return ''
        base=code.replace('_DIFF','')
        if part=='tel':
            if base in ['C01','C03']:
                s=self.m.val('tel',base+'_SEL')
                return '중' if s=='1)' else '하' if s=='2)' else '중'
            if base=='C05':
                s=self.m.val('tel','C05_SEL')
                return '상' if s=='1)' else '중'
            return '중'
        # CI fallback
        gp=self.m.val('ci','C01_SEL')
        gd='상' if gp=='1)' else '중' if gp=='2)' else '하' if gp=='3)' else '중'
        if base=='C01': return gd
        if base in ['C04','C06','C08','C10','C12']:
            s=self.m.val('ci',base+'_SEL')
            return '상' if s=='1)' and gd=='상' else '중' if s=='1)' else '하' if s=='2)' else '중'
        if base in ['C14','C23']:
            s=self.m.val('ci',base+'_SEL')
            return 'SPI-내부' if s=='1)' else 'SPI-외부' if s=='2)' else gd
        if base in ['C17','C18','C26','C27','C28','C32']:
            return gd
        if base in ['C19','C21']:
            s=self.m.val('ci',base+'_SEL')
            return '상' if s=='1)' and gd=='상' else '중' if s=='1)' else '하' if s=='2)' else '중'
        if base=='C29':
            s=self.m.val('ci','C29_SEL')
            return '상' if s=='1)' else '중' if s=='2)' else '하' if s=='3)' else '중'
        return gd

    def wrap_select_text(self, text, line_chars=24, max_lines=3):
        """Return 2~3 line text for long select display/menu labels."""
        t=str(text or '').replace('\n',' ').strip()
        if len(t) <= line_chars:
            return t
        words=t.split(' ')
        lines=[]
        cur=''
        for w in words:
            if not cur:
                cur=w
            elif len(cur)+1+len(w) <= line_chars:
                cur += ' ' + w
            else:
                lines.append(cur)
                cur=w
                if len(lines) >= max_lines-1:
                    break
        rest=' '.join(words[words.index(w):]) if False else ''
        if cur:
            lines.append(cur)
        # If text remains beyond max_lines, compose the final line from the remaining text and add ellipsis.
        joined=' '.join(lines)
        if len(joined) < len(t) and len(lines) >= max_lines:
            prefix=' '.join(lines[:-1])
            used=len(prefix) + (1 if prefix else 0)
            last=t[used:used+line_chars-1].strip() + '…'
            lines=lines[:-1]+[last]
        elif len(lines) > max_lines:
            lines=lines[:max_lines]
            lines[-1]=lines[-1][:line_chars-1].rstrip()+'…'
        return '\n'.join(lines)

    def set_select_value(self, rec, var, raw_value, display_text):
        var.set(self.wrap_select_text(display_text))
        self.update_input(rec, raw_value)

    def render_input_editor(self):
        for w in self.input_frame.inner.winfo_children(): w.destroy()
        heads=['Part','Code','구분','중분류','입력 항목','입력값','단위','Note']; widths=[13,13,12,34,50,32,18,72]
        for j,h in enumerate(heads): tk.Label(self.input_frame.inner,text=h,bg=HEAD,fg='white',font=('Segoe UI',9,'bold'),padx=6,pady=6).grid(row=0,column=j,sticky='nsew',padx=1,pady=1)
        row=1
        for x in self.m.inputs:
            if x.get('code')=='CTRL_EXTERNAL_MIN': x['value']=self.m.external_min
            if self.m.part!='both' and x['part']!=self.m.part: continue
            kind='자동결과' if x['type']=='computed' else 'Master연동' if x['code']=='CTRL_EXTERNAL_MIN' else '필수' if x['code'].startswith('CTRL') or x['code'].endswith('_QTY') or x['code'].endswith('_VALUE') else '선택'
            bg=GREEN if kind=='자동결과' else '#DDEBF7' if kind=='Master연동' else YELLOW if kind=='필수' else PURPLE
            vals=[self.part_label(x['part']),x['code'],kind,x.get('middle',''),x.get('detail','')]
            for j,v in enumerate(vals): tk.Label(self.input_frame.inner,text=v,bg=CARD if j!=2 else bg,fg=TEXT,font=('Segoe UI',9),wraplength=widths[j]*8,anchor='w',justify='left',padx=5,pady=4).grid(row=row,column=j,sticky='nsew',padx=1,pady=1)
            display_value=x.get('value','')
            if x.get('type')=='computed':
                display_value=self.fallback_diff(x.get('part'),x.get('code'))
                x['value']=display_value
            if x.get('code')=='CTRL_EXTERNAL_MIN':
                display_value=self.m.external_min
                x['value']=display_value
            var=tk.StringVar(value=str(display_value))
            if x['type']=='computed':
                # 자동결과는 Entry 대신 Label로 강제 표시해서 빈칸 문제가 없게 함
                tk.Label(self.input_frame.inner,text=str(display_value),bg=GREEN,fg=TEXT,font=('Segoe UI',9,'bold'),anchor='w',padx=6,pady=4).grid(row=row,column=5,sticky='nsew',padx=1,pady=1)
            elif x['code']=='CTRL_EXTERNAL_MIN':
                # 외주최소화는 왼쪽 Master Control 값 그대로 표시. 수정/선택 불가.
                master_val=self.m.external_min
                x['value']=master_val
                tk.Label(self.input_frame.inner,text=str(master_val),bg=YELLOW,fg=TEXT,font=('Segoe UI',9,'bold'),anchor='w',padx=6,pady=4).grid(row=row,column=5,sticky='nsew',padx=1,pady=1)
            elif x['type']=='select':
                display={o['text']:o['value'] for o in x['options']}
                rev={o['value']:o['text'] for o in x['options']}
                full_text=rev.get(x.get('value'),x.get('value',''))
                var.set(self.wrap_select_text(full_text))
                mb=tk.Menubutton(self.input_frame.inner,textvariable=var,relief='solid',bd=1,bg='white',fg=TEXT,activebackground='#DDEBF7',font=('Segoe UI',9),anchor='w',justify='left',wraplength=widths[5]*8-18,padx=6,pady=5)
                menu=tk.Menu(mb,tearoff=0,bg='white',fg=TEXT,activebackground='#DDEBF7',font=('Segoe UI',9))
                for txt,raw in display.items():
                    menu.add_command(label=self.wrap_select_text(txt),command=lambda rec=x,v=var,rv=raw,dt=txt:self.set_select_value(rec,v,rv,dt))
                mb.configure(menu=menu)
                mb.grid(row=row,column=5,sticky='ew',padx=1,pady=1)
            else:
                ent=ttk.Entry(self.input_frame.inner,textvariable=var,width=widths[5]); ent.grid(row=row,column=5,sticky='ew',padx=1,pady=1); ent.bind('<Return>',lambda e,rec=x,v=var:self.update_input(rec,v.get())); ent.bind('<FocusOut>',lambda e,rec=x,v=var:self.update_input(rec,v.get()))
            tk.Label(self.input_frame.inner,text=x.get('unit',''),bg=CARD,font=('Segoe UI',9)).grid(row=row,column=6,sticky='nsew',padx=1,pady=1); note_text=x.get('note','')
            if x['code']=='CTRL_EXTERNAL_MIN': note_text='왼쪽 Master Control 값이 그대로 표시됩니다. 이 행에서는 선택/수정하지 않습니다.'
            tk.Label(self.input_frame.inner,text=note_text,bg=CARD,fg='#506980',font=('Segoe UI',8),wraplength=460,anchor='w',justify='left').grid(row=row,column=7,sticky='nsew',padx=1,pady=1)
            row+=1
        for c,w in enumerate(widths): self.input_frame.inner.columnconfigure(c,minsize=w*8)
    def update_input(self,rec,value):
        if rec.get('value')==value: return
        self.m.input_changed(rec,value); self.refresh_all(); self.save_user_state(silent=True)
    def clear(self,t): t.delete(*t.get_children())
    def render_summary(self):
        area=self.summary_area
        for w in area.winfo_children(): w.destroy()
        m=self.m
        # calculate all-part totals independent of current part filter, then restore
        old_part=m.part
        def calc_for(part):
            m.part=part
            h,e,t=m.summary()
            mm=t/m.base_mh if m.base_mh else 0
            avg=mm/m.months if m.months else 0
            return h,e,t,mm,avg
        ci=calc_for('ci')
        tel=calc_for('tel')
        m.part='both'
        allv=calc_for('both')
        m.part=old_part
        ip,ep=m.pct()
        case_label='외주-단종 Case' if self.vars['case'].get()=='외주-단종' else '외주-종합 Case'
        rate=m.rate_short if self.vars['case'].get()=='외주-단종' else m.rate_comp
        # visual helpers
        def cell(parent,r,c,text,cs=1,rs=1,bg='white',fg=TEXT,bold=False,size=10,anchor='center',width=14,pady=6):
            lab=tk.Label(parent,text=text,bg=bg,fg=fg,font=('Segoe UI',size,'bold' if bold else 'normal'),bd=1,relief='solid',padx=8,pady=pady,anchor=anchor,justify='center',wraplength=max(80,width*9))
            lab.grid(row=r,column=c,columnspan=cs,rowspan=rs,sticky='nsew')
            return lab
        title=tk.Label(area,text='FEED사업 M/H Summary (C&I + 통신)',bg='white',fg='black',font=('Segoe UI',20,'bold'),pady=10)
        title.grid(row=0,column=0,columnspan=10,sticky='ew')
        proj=self.vars.get('project').get() if 'project' in self.vars else ''
        cell(area,1,0,'PROJECT:',cs=2,bg='white',bold=True,anchor='w')
        cell(area,1,2,proj or '-',cs=5,bg='white',anchor='w')
        cell(area,1,7,'Master 외주최소화',bg='#F7FAFD',bold=True)
        cell(area,1,8,self.m.external_min,bg=YELLOW if self.m.external_min=='Yes' else '#F7FAFD',bold=True)
        cell(area,1,9,f'내부/외부 {ip:.1f}/{ep:.1f}',bg='#F7FAFD',bold=True)
        cell(area,2,0,'1. M/H Summary',cs=10,bg='white',bold=True,size=12,anchor='w')
        # top summary table
        cell(area,3,0,'구분',rs=2,bg='#F8F8F8',bold=True,width=12)
        cell(area,3,1,'합계',cs=4,bg='#F8F8F8',bold=True,width=14)
        cell(area,4,1,'내부',bg='#F8F8F8',bold=True); cell(area,4,2,'외주',bg='#F8F8F8',bold=True); cell(area,4,3,'Total',bg='#F8F8F8',bold=True); cell(area,4,4,'예가',bg='#F8F8F8',bold=True)
        cell(area,5,0,'M/H',bg='white',bold=True); cell(area,5,1,fmt(allv[0])); cell(area,5,2,fmt(allv[1])); cell(area,5,3,fmt(allv[2])); cell(area,5,4,fmt(allv[1]*rate))
        cell(area,6,0,'M/M',bg='white',bold=True); cell(area,6,1,f1(allv[0]/m.base_mh if m.base_mh else 0)); cell(area,6,2,f1(allv[1]/m.base_mh if m.base_mh else 0)); cell(area,6,3,f1(allv[3])); cell(area,6,4,'-')
        cell(area,7,0,'투입 인원',bg='white',bold=True); cell(area,7,1,f1((allv[0]/m.base_mh/m.months) if m.base_mh and m.months else 0)); cell(area,7,2,f1((allv[1]/m.base_mh/m.months) if m.base_mh and m.months else 0)); cell(area,7,3,f1(allv[4])); cell(area,7,4,'-')
        cell(area,9,0,'2. M/H 상세내역',cs=10,bg='white',bold=True,size=12,anchor='w')
        # detail table similar to Excel
        cell(area,10,0,'구분',rs=3,bg='#F8F8F8',bold=True)
        cell(area,10,1,f'C&I설계 Part ({case_label})',cs=4,bg='#F8F8F8',bold=True,size=11)
        cell(area,10,5,f'통신설계 Part ({case_label})',cs=4,bg='#F8F8F8',bold=True,size=11)
        cell(area,11,1,'외주최소화 Ver.' if m.external_min=='Yes' else '일반 Ver.',cs=4,bg='white',fg='red',bold=True)
        cell(area,11,5,'외주최소화 Ver.' if m.external_min=='Yes' else '일반 Ver.',cs=4,bg='white',fg='red',bold=True)
        for c,text in enumerate(['내부','외주','Total','예가'],1): cell(area,12,c,text,bg='#F8F8F8',bold=True)
        for c,text in enumerate(['내부','외주','Total','예가'],5): cell(area,12,c,text,bg='#F8F8F8',bold=True)
        vals=[('M/H',0),('M/M',3),('투입 인원',4)]
        for rr,(label,idx) in enumerate(vals,13):
            cell(area,rr,0,label,bg='white',bold=True)
            if label=='M/H':
                cirow=[fmt(ci[0]),fmt(ci[1]),fmt(ci[2]),fmt(ci[1]*rate)]; telrow=[fmt(tel[0]),fmt(tel[1]),fmt(tel[2]),fmt(tel[1]*rate)]
            elif label=='M/M':
                cirow=[f1(ci[0]/m.base_mh if m.base_mh else 0),f1(ci[1]/m.base_mh if m.base_mh else 0),f1(ci[3]),'-']; telrow=[f1(tel[0]/m.base_mh if m.base_mh else 0),f1(tel[1]/m.base_mh if m.base_mh else 0),f1(tel[3]),'-']
            else:
                cirow=[f1((ci[0]/m.base_mh/m.months) if m.base_mh and m.months else 0),f1((ci[1]/m.base_mh/m.months) if m.base_mh and m.months else 0),f1(ci[4]),'-']; telrow=[f1((tel[0]/m.base_mh/m.months) if m.base_mh and m.months else 0),f1((tel[1]/m.base_mh/m.months) if m.base_mh and m.months else 0),f1(tel[4]),'-']
            for c,v in enumerate(cirow,1): cell(area,rr,c,v)
            for c,v in enumerate(telrow,5): cell(area,rr,c,v)
        cell(area,17,0,'Note',bg='#F7FAFD',bold=True)
        cell(area,17,1,'Summary는 왼쪽 Master Control, Input 수정값, 산출기준 Unit M/H를 기준으로 즉시 재계산됩니다.',cs=8,bg='#F7FAFD',anchor='w')
        for c in range(10): area.columnconfigure(c,weight=1,minsize=105)

    def render_output_form(self, part):
        sf=self.output_ci_frame if part=='ci' else self.output_tel_frame
        area=sf.inner
        for w in area.winfo_children(): w.destroy()
        m=self.m
        old_part=m.part
        m.part=part
        rows=[m.row(o) for o in m.visible_outputs()]
        h=sum(r['hec'] for r in rows); e=sum(r['ext'] for r in rows); total=sum(r['total'] for r in rows)
        mm=total/m.base_mh if m.base_mh else 0
        avg=mm/m.months if m.months else 0
        m.part=old_part
        title='C&I설계 화공 FEED사업 M/H 산출서 (Rev.3)' if part=='ci' else '통신설계 화공 FEED사업 M/H 산출서 (Rev.3)'
        ver='[ 외주최소화 Ver. ]' if m.external_min=='Yes' else '[ 일반 Ver. ]'
        part_name='CI (계장)' if part=='ci' else 'TEL (통신)'
        def cell(r,c,text,cs=1,rs=1,bg='white',fg=TEXT,bold=False,size=9,anchor='center',wrap=160,padx=5,pady=5):
            lab=tk.Label(area,text=text,bg=bg,fg=fg,font=('Segoe UI',size,'bold' if bold else 'normal'),bd=1,relief='solid',padx=padx,pady=pady,anchor=anchor,justify='center',wraplength=wrap)
            lab.grid(row=r,column=c,columnspan=cs,rowspan=rs,sticky='nsew')
            return lab
        # header
        cell(0,0,title,cs=11,bg='white',fg='black',bold=True,size=18,pady=8,wrap=900)
        cell(1,0,ver,cs=11,bg='white',fg='red',bold=True,size=15,pady=6,wrap=400)
        # base summary box
        cell(3,0,'산출 Base Data Summary',cs=4,bg='white',bold=True,anchor='w',wrap=300)
        if part=='ci':
            base_items=[('Total ICSS (DCS+ESD+F&G)+MMS IO 수량',m.nval('ci','A01_QTY'),'Point'),('Total Instrument 수량',m.nval('ci','A02_QTY'),'Set')]
        else:
            base_items=[('Telecom System 수량',m.nval('tel','A01_QTY'),'Set'),('Site Area',m.nval('tel','A02_QTY'),'m2'),('Building 수량',m.nval('tel','A03_QTY'),'Ea')]
        br=4
        for label,val,unit in base_items:
            cell(br,0,label,cs=2,bg='white',bold=True,anchor='w',wrap=320)
            cell(br,2,fmt(val),bg='#47D764',fg='black',bold=True)
            cell(br,3,unit,bg='white',fg='red',bold=True)
            br+=1
        cell(3,8,'Notes',cs=3,bg='#D9D9D9',bold=True,anchor='w')
        note='왼쪽 Master Control의 외주최소화 값(No/Yes)이 Output 전체에 그대로 자동 반영됩니다.'
        cell(4,8,note,cs=3,rs=2,bg='#EEEEEE',anchor='w',wrap=330)
        cell(br+1,0,'PROJECT:',cs=2,bg='#47D764',bold=True,anchor='w')
        cell(br+1,2,self.vars.get('project').get() if 'project' in self.vars else '-',cs=9,bg='#9AD780',anchor='w')
        # table title
        start=br+3
        cell(start,0,'MAN-HOUR 산출 TABLE',cs=11,bg='#93D47C',bold=True,size=11,wrap=700)
        start+=1
        # group headers
        cell(start,0,'No.',rs=2,bg='#D8E8F8',bold=True)
        cell(start,1,'Activity',rs=2,bg='#D8E8F8',bold=True,wrap=260)
        cell(start,2,'단위',rs=2,bg='#D8E8F8',bold=True)
        cell(start,3,'수량',rs=2,bg='#D8E8F8',bold=True)
        cell(start,4,'난이도',rs=2,bg='#D8E8F8',bold=True)
        cell(start,5,'내부 M/H',cs=2,bg='#BFE7F3',bold=True)
        cell(start,7,'외부 M/H',cs=2,bg='#C6EFCE',bold=True)
        cell(start,9,'Total',cs=2,bg='#FFF2CC',bold=True)
        cell(start+1,5,'Unit M/H',bg='#BFE7F3',bold=True); cell(start+1,6,'M/H',bg='#BFE7F3',bold=True)
        cell(start+1,7,'Unit M/H',bg='#C6EFCE',bold=True); cell(start+1,8,'M/H',bg='#C6EFCE',bold=True)
        cell(start+1,9,'M/H',bg='#FFF2CC',bold=True); cell(start+1,10,'비고',bg='#FFF2CC',bold=True)
        r=start+2
        for idx,row in enumerate(rows,1):
            bg='#FFFFFF' if idx%2 else '#F7FAFD'
            cell(r,0,idx,bg=bg); cell(r,1,row['activity'],bg=bg,anchor='w',wrap=330); cell(r,2,row['unit'],bg=bg); cell(r,3,f1(row['qty']),bg=bg); cell(r,4,row['diff'],bg=bg)
            cell(r,5,f1(row['iu']),bg=bg); cell(r,6,fmt(row['hec']),bg=bg); cell(r,7,f1(row['eu']),bg=bg); cell(r,8,fmt(row['ext']),bg=bg); cell(r,9,fmt(row['total']),bg=bg); cell(r,10,'',bg=bg)
            r+=1
        # totals band
        cell(r,0,'Base M/H (1 M/M)',cs=2,bg='yellow',bold=True,anchor='w'); cell(r,2,fmt(m.base_mh),bg='yellow',bold=True); cell(r,3,'M/H',bg='yellow'); cell(r,4,'합계 M/H',bg='yellow',bold=True); cell(r,5,fmt(h),cs=2,bg='yellow',bold=True); cell(r,7,fmt(e),cs=2,bg='yellow',bold=True); cell(r,9,fmt(total),cs=2,bg='yellow',bold=True); r+=1
        cell(r,0,'설계기간',cs=2,bg='white',bold=True,anchor='w'); cell(r,2,f1(m.months),bg='white',bold=True); cell(r,3,'개월',bg='white'); cell(r,4,'합계 M/M',bg='white',bold=True); cell(r,5,f1(h/m.base_mh if m.base_mh else 0),cs=2,bg='white',bold=True); cell(r,7,f1(e/m.base_mh if m.base_mh else 0),cs=2,bg='white',bold=True); cell(r,9,f1(mm),cs=2,bg='white',bold=True); r+=1
        rate=m.rate_short if self.vars.get('case').get()=='외주-단종' else m.rate_comp
        cell(r,0,'외주 적용단가',cs=2,bg='white',bold=True,anchor='w'); cell(r,2,fmt(rate),bg='white',bold=True); cell(r,3,'원',bg='white'); cell(r,4,'평균 투입 인원',bg='white',bold=True); cell(r,5,f1((h/m.base_mh/m.months) if m.base_mh and m.months else 0),cs=2,bg='white',bold=True); cell(r,7,f1((e/m.base_mh/m.months) if m.base_mh and m.months else 0),cs=2,bg='white',bold=True); cell(r,9,f1(avg),cs=2,bg='white',bold=True); r+=1
        cell(r,0,'외주 예가',cs=4,bg='#FFF2CC',bold=True,anchor='w'); cell(r,4,fmt(e*rate),cs=7,bg='#FFF2CC',bold=True,fg='red')
        # column sizes
        widths=[42,310,70,80,70,90,90,90,90,90,150]
        for c,w in enumerate(widths): area.columnconfigure(c,minsize=w,weight=1 if c==1 else 0)

    def render_op1(self):
        if hasattr(self,'op_grid_op1'):
            self.op_grid_op1.render()

    def render_op2(self,name):
        if name=='OP2-단종' and hasattr(self,'op_grid_short'):
            self.op_grid_short.render()
        elif name=='OP2-종합' and hasattr(self,'op_grid_comp'):
            self.op_grid_comp.render()

    def render_std(self):
        if hasattr(self,'std_grid_ci'): self.std_grid_ci.render()
        if hasattr(self,'std_grid_tel'): self.std_grid_tel.render()


    def update_std_value(self, part, key, typ, diff, value):
        """Update editable 산출기준 value and recalculate all output tabs."""
        try:
            self.m.std[part][key].setdefault(typ,{})[diff]=value
            target_activity=self.m.std[part][key].get('activity')
            for o in self.m.outputs:
                if o.get('part')==part and isinstance(o.get('std'),dict) and o['std'].get('activity')==target_activity:
                    o['std'].setdefault(typ,{})[diff]=value
            self.m.recalc()
            self.render_summary(); self.render_output_form('ci'); self.render_output_form('tel')
            self.render_op1(); self.render_op2('OP2-단종'); self.render_op2('OP2-종합'); self.render_std()
            self.save_user_state(silent=True)
        except Exception as e:
            messagebox.showerror('산출기준 수정 실패', str(e))

    def on_tree_double_click(self, event):
        """Edit 산출기준 numeric Unit M/H values directly in the Treeview."""
        tree=event.widget
        # only standards tabs are editable
        if tree not in (self.trees.get('산출기준_CI'), self.trees.get('산출기준_TEL')):
            return
        item=tree.identify_row(event.y); col_id=tree.identify_column(event.x)
        if not item or item not in self.std_item_map:
            return
        try:
            col=int(col_id.replace('#',''))-1
        except Exception:
            return
        part,key=self.std_item_map[item]
        # Define editable columns and target field
        if part=='ci':
            mapping={2:('int','SPI'),3:('int','상'),4:('int','중'),5:('int','하'),6:('ext','SPI'),7:('ext','상'),8:('ext','중'),9:('ext','하')}
        else:
            mapping={2:('int','상'),3:('int','중'),4:('int','하'),5:('ext','상'),6:('ext','중'),7:('ext','하')}
        if col not in mapping:
            return
        bbox=tree.bbox(item, col_id)
        if not bbox:
            return
        x,y,w,h=bbox
        old=tree.set(item, col_id)
        editor=ttk.Entry(tree)
        editor.place(x=x, y=y, width=w, height=h)
        editor.insert(0, old)
        editor.focus_set()
        editor.select_range(0, 'end')
        def commit(event=None):
            val=editor.get().strip()
            try:
                n=float(val.replace(',','')) if val!='' else 0.0
            except Exception:
                messagebox.showerror('입력 오류','숫자만 입력할 수 있습니다.')
                editor.focus_set(); return
            typ,diff=mapping[col]
            # update master standard table
            st=self.m.std[part][key]
            st.setdefault(typ,{})[diff]=n
            # update output rows that reference same standard activity
            target_activity=st.get('activity')
            for o in self.m.outputs:
                if o.get('part')==part and isinstance(o.get('std'),dict):
                    if o['std'].get('activity')==target_activity:
                        o['std'].setdefault(typ,{})[diff]=n
            editor.destroy()
            self.m.recalc()
            self.render_summary(); self.render_output_form('ci'); self.render_output_form('tel'); self.render_op1(); self.render_op2('OP2-단종'); self.render_op2('OP2-종합'); self.render_std()
            self.save_user_state(silent=True)
        def cancel(event=None):
            try: editor.destroy()
            except Exception: pass
        editor.bind('<Return>', commit)
        editor.bind('<FocusOut>', commit)
        editor.bind('<Escape>', cancel)

    def _make_logo_widget(self, parent):
        # Show embedded Hyundai Engineering logo in program UI.
        try:
            logo_path=Path(tempfile.gettempdir())/'hyundai_engineering_logo_program.png'
            logo_path.write_bytes(base64.b64decode(HYUNDAI_LOGO_B64))
            self.logo_img=tk.PhotoImage(file=str(logo_path))
            # keep original logo size from uploaded file; image is intentionally compact
            lbl=tk.Label(parent,image=self.logo_img,bg=CARD,bd=0)
            lbl.pack(anchor='w',pady=(0,10))
        except Exception:
            # If image rendering fails, keep text fallback but do not block program use.
            ttk.Label(parent,text='HYUNDAI ENGINEERING',style='Card.TLabel',font=('Segoe UI',12,'bold'),foreground=PRIMARY).pack(anchor='w',pady=(0,10))

    def _default_report_dir(self):
        # Report 저장 기본 위치: 현재 사용자 문서 폴더
        candidates = [Path.home()/'Documents', Path.home()/'문서', Path.home()/'OneDrive'/'Documents', Path.home()]
        for p in candidates:
            try:
                if p.exists() and p.is_dir():
                    return p
            except Exception:
                pass
        return Path.home()
    def _open_folder(self, folder):
        try:
            if sys.platform.startswith('win'):
                os.startfile(str(folder))
            elif sys.platform == 'darwin':
                subprocess.Popen(['open', str(folder)])
            else:
                subprocess.Popen(['xdg-open', str(folder)])
        except Exception:
            pass

    def _report_data(self):
        self.apply_left()
        m=self.m
        old=m.part
        def calc(part):
            m.part=part
            rows=[m.row(o) for o in m.visible_outputs()]
            h=sum(r['hec'] for r in rows); e=sum(r['ext'] for r in rows); t=sum(r['total'] for r in rows)
            return {'rows':rows,'internal':h,'external':e,'total':t,'mm':t/m.base_mh if m.base_mh else 0,'avg':(t/m.base_mh/m.months) if m.base_mh and m.months else 0}
        ci=calc('ci'); tel=calc('tel')
        total={'internal':ci['internal']+tel['internal'],'external':ci['external']+tel['external'],'total':ci['total']+tel['total']}
        total['mm']=total['total']/m.base_mh if m.base_mh else 0
        total['avg']=total['mm']/m.months if m.months else 0
        m.part=old
        return {'project':self.vars.get('project').get() if 'project' in self.vars else '-', 'external_min':m.external_min, 'base_mh':m.base_mh, 'months':m.months, 'rate_short':m.rate_short, 'rate_comp':m.rate_comp, 'ci':ci, 'tel':tel, 'total':total, 'short_cost':total['external']*m.rate_short, 'comp_cost':total['external']*m.rate_comp, 'std':m.std}
    def state_file(self):
        return Path(__file__).with_name('FEED_MH_Calculator_Last_Input.json')

    def collect_user_state(self):
        return {
            'project': self.m.project,
            'part': self.m.part,
            'case': self.m.case,
            'ratio_mode': self.m.ratio_mode,
            'external_min': self.m.external_min,
            'base_mh': self.m.base_mh,
            'months': self.m.months,
            'rate_short': self.m.rate_short,
            'rate_comp': self.m.rate_comp,
            'internal_pct': self.m.internal_pct,
            'external_pct': self.m.external_pct,
            'inputs': [{'part': x.get('part'), 'code': x.get('code'), 'value': x.get('value')} for x in self.m.inputs],
            'std': self.m.std
        }

    def apply_user_state_to_vars(self):
        self.vars['part'].set(P.get(self.m.part,'전체'))
        self.vars['case'].set('외주-단종' if self.m.case=='short' else '외주-종합')
        self.vars['ratio'].set('원본 기준' if self.m.ratio_mode=='original' else '사용자 입력')
        self.vars['external'].set(self.m.external_min)
        self.vars['project'].set(str(self.m.project))
        self.vars['base'].set(str(self.m.base_mh))
        self.vars['months'].set(str(self.m.months))
        self.vars['rate_short'].set(str(self.m.rate_short))
        self.vars['rate_comp'].set(str(self.m.rate_comp))
        self.vars['ip'].set(f1(self.m.internal_pct))
        self.vars['ep'].set(f1(self.m.external_pct))

    def save_user_state(self, silent=False):
        try:
            self.state_file().write_text(json.dumps(self.collect_user_state(), ensure_ascii=False, indent=2), encoding='utf-8')
            if not silent:
                messagebox.showinfo('저장 완료', f'현재 입력값을 저장했습니다.\n{self.state_file()}')
        except Exception as e:
            if not silent:
                messagebox.showerror('저장 실패', str(e))

    def load_user_state(self, show_msg=False):
        p = self.state_file()
        if not p.exists():
            if show_msg:
                messagebox.showinfo('저장값 없음', '저장된 입력값이 없습니다.')
            return
        try:
            data = json.loads(p.read_text(encoding='utf-8'))
            self.m.project = data.get('project','')
            self.m.part = data.get('part','both')
            self.m.case = data.get('case','short')
            self.m.ratio_mode = data.get('ratio_mode','original')
            self.m.external_min = data.get('external_min','No')
            self.m.base_mh = num(data.get('base_mh',161))
            self.m.months = num(data.get('months',4))
            self.m.rate_short = num(data.get('rate_short',32000))
            self.m.rate_comp = num(data.get('rate_comp',39000))
            self.m.internal_pct = num(data.get('internal_pct',0))
            self.m.external_pct = num(data.get('external_pct',0))
            mp = {(x.get('part'), x.get('code')): x for x in self.m.inputs}
            for item in data.get('inputs',[]):
                key=(item.get('part'), item.get('code'))
                if key in mp:
                    mp[key]['value'] = item.get('value')
            # 산출기준 Unit M/H 사용자 조정값 복원
            std_data=data.get('std')
            if isinstance(std_data, dict):
                for part in ['ci','tel']:
                    if isinstance(std_data.get(part), dict):
                        for sk, sv in std_data[part].items():
                            if sk in self.m.std[part] and isinstance(sv, dict):
                                self.m.std[part][sk].update(sv)
                # 출력 행 안에 복사되어 있는 std도 같이 업데이트
                for o in self.m.outputs:
                    if not isinstance(o.get('std'), dict):
                        continue
                    part=o.get('part')
                    activity=o['std'].get('activity')
                    for st in self.m.std.get(part, {}).values():
                        if st.get('activity') == activity:
                            o['std'].update(st)
                            break
            self.m.sync_common(); self.m.recalc(); self.apply_user_state_to_vars()
            if show_msg:
                messagebox.showinfo('불러오기 완료', '저장된 입력값을 불러왔습니다.')
        except Exception as e:
            if show_msg:
                messagebox.showerror('불러오기 실패', str(e))

    def export_user_state(self):
        try:
            path=filedialog.asksaveasfilename(defaultextension='.json', filetypes=[('JSON 파일','*.json')], initialfile='FEED_MH_Calculator_Input_Backup.json')
            if not path:
                return
            Path(path).write_text(json.dumps(self.collect_user_state(), ensure_ascii=False, indent=2), encoding='utf-8')
            messagebox.showinfo('내보내기 완료', f'현재 입력값을 파일로 저장했습니다.\n{path}')
        except Exception as e:
            messagebox.showerror('내보내기 실패', str(e))

    def import_user_state(self):
        try:
            path=filedialog.askopenfilename(filetypes=[('JSON 파일','*.json'),('모든 파일','*.*')])
            if not path:
                return
            data=json.loads(Path(path).read_text(encoding='utf-8'))
            # 선택한 파일을 기본 저장 파일로 복사한 뒤 일반 불러오기 로직을 사용합니다.
            self.state_file().write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding='utf-8')
            self.load_user_state(show_msg=False)
            self.refresh_all()
            messagebox.showinfo('가져오기 완료', '선택한 입력값 파일을 불러왔습니다.')
        except Exception as e:
            messagebox.showerror('가져오기 실패', str(e))

    def reset_user_state(self):
        if not messagebox.askyesno('초기화 확인', '입력값을 초기화하시겠습니까?\n저장된 마지막 입력값도 삭제됩니다.'):
            return
        try:
            if self.state_file().exists():
                self.state_file().unlink()
            self.m = Model()
            self.apply_user_state_to_vars()
            self.refresh_all()
            messagebox.showinfo('초기화 완료', '입력값을 초기화했습니다.')
        except Exception as e:
            messagebox.showerror('초기화 실패', str(e))

    def on_close(self):
        self.save_user_state(silent=True)
        self.destroy()

    def generate_report(self):
        data=self._report_data()
        default_dir=self._default_report_dir()
        try:
            path=filedialog.asksaveasfilename(initialdir=str(default_dir), defaultextension='.docx',filetypes=[('Word Document','*.docx')],initialfile='FEED_MH_Report_Hyundai.docx')
            if path:
                report_utils.create_word_report(data, Path(path), HYUNDAI_LOGO_B64)
                if not Path(path).exists():
                    raise RuntimeError('Word 파일이 생성되지 않았습니다.')
                self._open_folder(Path(path).parent)
                messagebox.showinfo('완료',f'Word Report 생성 완료:\n{path}')
        except Exception as e:
            messagebox.showerror('Word Report 생성 실패', 'Word Report 생성 중 오류가 발생했습니다.\n\n가능 원인:\n1) 압축을 풀지 않고 ZIP 안에서 실행함\n2) report_utils.py 파일이 프로그램과 같은 폴더에 없음\n3) 저장 폴더 권한 문제\n\n해결:\nZIP 파일을 반드시 압축 해제한 후 FEED_MH_Standalone_Calculator_V0_1.py를 실행하세요.\n\n오류 내용:\n' + str(e))

if __name__=='__main__': App().mainloop()
